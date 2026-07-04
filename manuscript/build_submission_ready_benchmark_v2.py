from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT = ROOT / "outputs"
DOCX = OUT / "donor_disease_aware_sc_aging_submission_ready_v2.docx"
MD = ROOT / "manuscript" / "donor_disease_aware_sc_aging_submission_ready_v2.md"

meta_summary = pd.read_json(ROOT / "analysis/outputs/metadata_audit_summary.json", typ="series")
tissue = pd.read_csv(ROOT / "analysis/outputs/metadata_feasibility_map.tsv", sep="\t")
sens = pd.read_csv(ROOT / "analysis/outputs/threshold_sensitivity.tsv", sep="\t")
bench_audit = pd.read_json(ROOT / "analysis/benchmark_outputs/blood_tcell_embedding_benchmark_audit.json", typ="series")
bench = pd.read_csv(ROOT / "analysis/upgrade_outputs/blood_tcell_30seed_results_summary.tsv", sep="\t")
diag = pd.read_csv(ROOT / "analysis/upgrade_outputs/blood_tcell_30seed_leakage_diagnostics.tsv", sep="\t")
lodo = pd.read_csv(ROOT / "analysis/upgrade_outputs/lodo_dataset_holdout_metrics.tsv", sep="\t")
lodo_excl = pd.read_csv(ROOT / "analysis/upgrade_outputs/lodo_dataset_exclusions.tsv", sep="\t")
depth = pd.read_csv(ROOT / "analysis/upgrade_outputs/sampling_depth_sensitivity_summary.tsv", sep="\t")
extra = pd.read_csv(ROOT / "analysis/upgrade_outputs/extra_stratum_numeric_age_feasibility.tsv", sep="\t")


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def style_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_line_numbers(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = False
        style.paragraph_format.line_spacing = 1.15


def p(doc: Document, text: str):
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 2.0
    return para


def title_meta_paragraph(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.25
    return para


def add_superscript(para, text: str) -> None:
    run = para.add_run(text)
    run.font.superscript = True


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, df: pd.DataFrame, cols: list[str]) -> None:
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, col in enumerate(cols):
        set_cell(tbl.rows[0].cells[i], col, True)
        shade(tbl.rows[0].cells[i], "E8EEF5")
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(cols):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.3f}"
            set_cell(cells[i], str(value))
    doc.add_paragraph()


def pic(doc: Document, path: str, caption: str, width: float = 6.25) -> None:
    fp = ROOT / path
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.paragraph_format.line_spacing = 1.15
        cap.runs[0].italic = True


def ref(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.line_spacing = 1.0
    para.add_run(text)


def metric(model: str, split: str, col: str = "auroc_mean") -> float:
    row = bench[(bench["model"] == model) & (bench["split"] == split)]
    return float(row.iloc[0][col])


def metric_ci(model: str, split: str) -> str:
    row = bench[(bench["model"] == model) & (bench["split"] == split)].iloc[0]
    return f"{row['auroc_mean']:.3f} (95% CI {row['auroc_ci95_low']:.3f}-{row['auroc_ci95_high']:.3f})"


def depth_metric(model: str, split: str, d: int) -> float:
    row = depth[(depth["model"] == model) & (depth["split"] == split) & (depth["requested_cells_per_donor"] == d)]
    return float(row.iloc[0]["auroc_mean"])


def build_markdown() -> str:
    return f"""# Submission-ready benchmark draft v2

Updated empirical benchmark with 30 seeds, LODO diagnostics, confounding diagnosis, sampling-depth sensitivity, and reproducibility package.

Key result: blood T-cell dataset-holdout AUROC remained fragile for hosted cell-level embeddings, while disease-free donor-holdout donor-mean embeddings retained stronger signal.
"""


def build_docx() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Leakage-aware evaluation of CELLxGENE-hosted single-cell embeddings for blood T-cell aging prediction")
    r.bold = True
    r.font.size = Pt(16)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Computational benchmark manuscript for public CELLxGENE data").italic = True
    authors = title_meta_paragraph(doc)
    authors.add_run("Yi Zha")
    add_superscript(authors, "1")
    authors.add_run(", Da Lin")
    add_superscript(authors, "1")
    authors.add_run(", Ying Chen")
    add_superscript(authors, "2")
    authors.add_run(", Yue Liu")
    add_superscript(authors, "2")
    authors.add_run(", Yu Zhang")
    add_superscript(authors, "1,*")
    aff1 = title_meta_paragraph(doc)
    add_superscript(aff1, "1")
    aff1.add_run(" The Second Affiliated Hospital of Wenzhou Medical University, Wenzhou, Zhejiang, China.")
    aff2 = title_meta_paragraph(doc)
    add_superscript(aff2, "2")
    aff2.add_run(" Wenzhou Medical University, Wenzhou, Zhejiang, China.")
    corr = title_meta_paragraph(doc)
    corr.add_run("*Correspondence: Yu Zhang, zhangyu1@wzhealth.com; ORCID: 0000-0001-8579-3692.")

    doc.add_heading("Abstract", level=1)
    p(doc, "Background: Public single-cell atlases and hosted embeddings are increasingly reused for prediction tasks, but random-cell splits can overestimate generalizable performance when labels are structured by donor identity, disease status, assay, or study origin. Aging prediction from public metadata-derived labels is especially vulnerable to this problem because age groups may be unevenly distributed across datasets.")
    p(doc, f"Methods: We developed a leakage-aware computational benchmark for CELLxGENE-hosted public single-cell resources. After a metadata feasibility audit of CELLxGENE Census release {meta_summary['census_version']}, we completed the empirical benchmark in blood T cells. Young and old donors were defined using development-stage-derived numeric ages of <=40 and >=60 years. The benchmark used {int(bench_audit['sample_all_donors'])} balanced donors and {int(bench_audit['sample_all_cells'])} sampled cells, with a separate disease-free subset of {int(bench_audit['sample_disease_free_donors'])} donors and {int(bench_audit['sample_disease_free_cells'])} cells. We evaluated CELLxGENE-hosted scVI and TranscriptFormer TF-Sapiens embeddings, donor-mean embeddings, and a metadata-only baseline using 30 repeated seeds across random-cell, donor-holdout, grouped dataset-holdout, disease-free, label-shuffling, leave-one-dataset-out feasibility, confounding, and sampling-depth analyses.")
    p(doc, f"Results: Random-cell AUROC in blood T cells was {metric_ci('metadata_confounder_logistic','random_cell')} for metadata alone, {metric_ci('TranscriptFormer_embedding_logistic','random_cell')} for TranscriptFormer, and {metric_ci('scVI_embedding_logistic','random_cell')} for scVI, showing that metadata-only prediction exceeded or matched the hosted embedding models under the easiest split. Dataset-holdout AUROC was lower and more variable: {metric_ci('metadata_confounder_logistic','dataset_holdout')} for metadata, {metric_ci('TranscriptFormer_embedding_logistic','dataset_holdout')} for TranscriptFormer, and {metric_ci('scVI_embedding_logistic','dataset_holdout')} for scVI. Disease-free donor-holdout performance was stronger for donor-mean embeddings, reaching {metric_ci('TranscriptFormer_mean_embedding_pseudobulk','disease_free_donor_holdout')} for TranscriptFormer and {metric_ci('scVI_mean_embedding_pseudobulk','disease_free_donor_holdout')} for scVI. LODO evaluation was feasible for only one held-out dataset, showing limited cross-study validation capacity under the available public metadata.")
    p(doc, "Conclusions: This study provides a reproducible public-data evaluation framework for hosted single-cell embeddings and aging prediction. The blood T-cell benchmark shows that random-cell evaluation can overestimate generalizable performance and that donor-, disease-, and dataset-aware diagnostics are necessary before interpreting public single-cell age-prediction results biologically.")

    doc.add_heading("Keywords", level=1)
    p(doc, "single-cell RNA sequencing; aging-related prediction; CELLxGENE Census; hosted embeddings; scVI; TranscriptFormer; metadata leakage; dataset shift; benchmark")

    doc.add_heading("Background", level=1)
    p(doc, "Public single-cell atlases are increasingly reused as computational benchmark resources because they combine large cell numbers, harmonized metadata, and reusable representation layers. The Human Cell Atlas, Human Cell Landscape, and Tabula Sapiens established shared infrastructures for cross-tissue and cross-study single-cell analysis [14-16]. CZ CELLxGENE Discover extends this ecosystem by aggregating public datasets and exposing hosted analysis resources, including embeddings that can be used without retraining every model from raw expression matrices [17].")
    p(doc, "Aging-related prediction is an attractive test case for hosted single-cell representations. Molecular clocks showed that age can be estimated from high-dimensional molecular profiles, and single-cell aging atlases have shown that age-associated structure can be cell-type- and tissue-dependent [1,2,7]. Human single-cell studies have reported age-associated structure in skin, transcriptomic age estimates, and skeletal muscle [8-10]. Multi-tissue aging resources further show why age-related structure should be evaluated in a cell-type-aware manner [11,12]. However, public atlas labels are observational metadata rather than randomized experimental labels, so prediction performance can reflect source structure as well as biological variation.")
    p(doc, "This problem is especially important when age labels are derived from public metadata. Age groups can be unevenly distributed across donors, diseases, assays, studies, and recruitment settings. Best-practice guidance and integration methods show that preprocessing and latent representations can strongly shape downstream analyses [25-27]. Batch effects can dominate benchmark behavior if they are not explicitly diagnosed [28-30]. Atlas-scale dataset composition creates a related source of downstream instability [31]. A public-data benchmark therefore needs to evaluate metadata leakage and dataset shift, not only random-split accuracy.")
    p(doc, "Representation models add another reason for careful validation. scVI established a probabilistic latent-variable framework for single-cell expression [18,19]. Transcriptomic representation models such as Geneformer, scGPT, and scFoundation expanded the space of reusable cellular embeddings [20-22]. GeneCompass and TranscriptFormer further illustrate the rapid movement toward pretrained transcriptomic representations [23,24]. These models can encode useful biological structure, but flexible embeddings can also preserve study, assay, and cohort structure. The present study directly evaluates only CELLxGENE-hosted scVI and TranscriptFormer representations.")
    p(doc, "Previous single-cell aging atlases and representation-model studies have not, by themselves, provided a donor-, disease-, and study-aware benchmark for aging-related prediction from hosted public embeddings. In particular, a computational benchmark should include metadata-only baselines, donor-aware splitting, disease-free sensitivity analysis, grouped dataset-holdout evaluation, label-shuffling controls, LODO feasibility diagnostics, and sampling-depth sensitivity. These design choices follow broader biomedical machine-learning concerns about leakage and small-sample validation [39,40]. They also reflect reporting and validation guidance for prediction models and external validation [41-43]. Clinical AI evaluation guidance reinforces the same need for transparent validation and conservative interpretation [44,45]. Dataset-shift and causal-interpretation limits are especially important when public metadata define both labels and sources [46,47].")
    p(doc, "Here, we contribute a leakage-aware benchmark workflow for hosted single-cell representations in public aging-related prediction, using blood T cells as the completed empirical stratum. The goal is not to claim a pan-tissue human aging predictor, but to provide a reproducible public-data evaluation framework that tests whether apparent prediction behavior persists under donor-, disease-, and study-aware controls.")

    doc.add_heading("Methods", level=1)
    pic(doc, "outputs/upgrade_figures/figure0_benchmark_design_schematic.png", "Figure 1. Leakage-aware benchmark design. CELLxGENE Census public metadata and hosted scVI and TranscriptFormer embeddings were evaluated in a blood T-cell young-versus-old prediction task. The workflow compares random-cell, donor-holdout, grouped dataset-holdout, disease-free, metadata-only, label-shuffling, LODO, and sampling-depth diagnostics to assess donor leakage, disease confounding, study shift, source predictability, and donor-aggregation sensitivity.", width=6.3)
    doc.add_heading("Metadata feasibility audit", level=2)
    p(doc, f"The Phase 0 audit used CELLxGENE Census release {meta_summary['census_version']}. We queried primary human-cell metadata for T cells, monocytes, macrophages, endothelial cells, and fibroblasts. A tissue-cell stratum was considered eligible when it contained at least 30 donors, at least 10 young/adult donors, at least 10 old/elderly donors, at least two datasets, and disease-free donor retention of at least 0.50. Threshold sensitivity varied donor, age-proxy, and disease-free requirements.")
    p(doc, "For empirical age prediction, development_stage labels were parsed for numeric year-old labels and decade labels. Donors with inferred numeric ages <=40 years were classified as young, donors >=60 years as old, and middle-aged donors were excluded from binary classification.")
    doc.add_heading("Blood T-cell benchmark", level=2)
    p(doc, f"The blood T-cell benchmark used {int(bench_audit['sample_all_donors'])} donors balanced by age group and sampled up to eight cells per donor, yielding {int(bench_audit['sample_all_cells'])} cells. A disease-free subset was sampled independently with {int(bench_audit['sample_disease_free_donors'])} donors and {int(bench_audit['sample_disease_free_cells'])} cells. Hosted embeddings were retrieved for scVI and TranscriptFormer TF-Sapiens from the same Census release and cached locally in the project reproducibility directory.")
    p(doc, "Cell-level models used logistic regression with class weighting. The metadata-only baseline used dataset_id, assay, disease, and sex. Embedding models used either the cell-level hosted embedding or the donor-level mean embedding created by averaging sampled cells within each donor. This baseline was included to quantify how much aging-related prediction could be obtained from source metadata alone.")
    p(doc, "The main repeated-seed evaluation used seeds 1301 through 1330. The random-cell split stratified cells by the young-old label without preventing cells from the same donor from appearing in both training and test sets. The donor-holdout split grouped cells by donor_id so that all cells from a donor were assigned to either training or test. The grouped dataset-holdout split grouped cells by dataset_id to test performance under study-level shifts.")
    p(doc, "Disease-free analyses repeated the same split logic after restricting to donors annotated as normal or healthy. Label-shuffling controls permuted the training labels while preserving the original split structure and test labels. LODO and sampling-depth sensitivity were treated as diagnostic analyses and are described below.")
    doc.add_heading("Leave-one-dataset-out and diagnostic analyses", level=2)
    p(doc, "LODO evaluation held out one dataset at a time when the held-out dataset contained at least two young and two old donors and the remaining training data retained both classes. Datasets failing this rule were recorded with exclusion reasons. Confounding diagnostics summarized donor age labels by dataset, disease, assay, and sex; estimated coefficients from a metadata-only logistic model; and projected hosted embeddings by principal components. Sampling-depth sensitivity repeated donor-mean embedding evaluation after sampling 1, 2, 4, 8, or 16 cells per donor. The 16-cell point was capped by the existing eight-cell-per-donor benchmark sample.")
    doc.add_heading("Statistical analysis", level=2)
    p(doc, "Primary metrics were AUROC, AUPRC, and balanced accuracy. Repeated-seed results are reported as mean, standard deviation, and Wald-style 95% confidence interval across 30 split seeds. These intervals quantify repeated split variability rather than uncertainty over independent external cohorts. Paired comparisons used seed-matched AUROC differences. Generalization drop was defined as random-cell AUROC minus donor-holdout AUROC. Permutation gap was defined as true-label random-cell AUROC minus shuffled-label random-cell AUROC.")
    doc.add_heading("Benchmark risks and diagnostic controls", level=2)
    p(doc, "The benchmark was organized around explicit leakage and confounding risks. Donor leakage was assessed by donor-holdout splits. Disease confounding was assessed by disease-free sensitivity analysis and disease metadata diagnostics. Study or dataset shift was assessed by grouped dataset-holdout and LODO feasibility diagnostics. Metadata-only predictability was assessed by the metadata-only baseline. Spurious label structure was assessed by label-shuffling controls. Donor aggregation sensitivity was assessed by sampling-depth sensitivity for donor-mean embeddings.")

    doc.add_heading("Results", level=1)
    doc.add_heading("Question 1: which public stratum supports a reproducible blood T-cell benchmark?", level=2)
    p(doc, f"The metadata audit retrieved {int(meta_summary['target_primary_cell_rows']):,} target primary-cell rows from {int(meta_summary['unique_donors']):,} donors, {int(meta_summary['unique_datasets']):,} datasets, {int(meta_summary['unique_collections']):,} collections, and {int(meta_summary['unique_tissues']):,} tissues. Eighteen tissue-cell strata met the primary proxy-based eligibility threshold, whereas 98 met pilot-candidate criteria.")
    pic(doc, "outputs/figures_v2/figure1_metadata_screening_flow.png", "Figure 2. Metadata screening flow for the CELLxGENE Census feasibility audit.")
    pic(doc, "outputs/figures_v2/figure4_disease_free_ci.png", "Figure 3. Disease-free donor retention by target cell system with 95% Wilson confidence intervals.")
    base = sens[(sens["donor_min"] == 30) & (sens["age_group_min"] == 10) & (sens["disease_free_min"] == 0.5)].iloc[0]
    strict = sens[(sens["donor_min"] == 40) & (sens["age_group_min"] == 10) & (sens["disease_free_min"] == 0.6)].iloc[0]
    p(doc, f"Threshold sensitivity showed that eligibility depended on the donor and disease-free criteria. The primary threshold yielded {int(base['eligible_strata_n'])} eligible strata, whereas the stricter rule requiring at least 40 donors and disease-free retention of 0.60 yielded {int(strict['eligible_strata_n'])} eligible strata.")
    pic(doc, "outputs/figures_v2/figure3_threshold_sensitivity.png", "Figure 4. Threshold sensitivity analysis across donor and disease-free retention thresholds.")

    doc.add_heading("Question 2: how much apparent aging-related signal is observed under random-cell splitting?", level=2)
    p(doc, f"In blood T cells, random-cell AUROC was {metric_ci('metadata_confounder_logistic','random_cell')} for metadata alone, {metric_ci('TranscriptFormer_embedding_logistic','random_cell')} for TranscriptFormer, and {metric_ci('scVI_embedding_logistic','random_cell')} for scVI. Metadata-only random-cell performance therefore exceeded TranscriptFormer and scVI cell-level performance, showing that dataset, assay, disease, and sex carried substantial age-predictive structure in the sampled public data.")
    pic(doc, "outputs/upgrade_figures/figure4_blood_tcell_30seed_performance.png", "Figure 5. Thirty-seed blood T-cell benchmark performance. AUROC values are shown for metadata-only, cell-level hosted embedding, and donor-mean embedding models across random-cell, donor-holdout, grouped dataset-holdout, and disease-free donor-holdout splits. The random-cell split permits cells from the same donor to appear across train and test, whereas donor-holdout and dataset-holdout splits evaluate donor-level and study-level transfer.", width=6.3)

    doc.add_heading("Question 3: how does performance change under donor-aware and dataset-aware evaluation?", level=2)
    p(doc, f"Donor-holdout AUROC was lower than or comparable to random-cell performance for cell-level embeddings, reaching {metric_ci('TranscriptFormer_embedding_logistic','donor_holdout')} for TranscriptFormer and {metric_ci('scVI_embedding_logistic','donor_holdout')} for scVI. Donor-level mean embeddings improved donor-holdout performance to {metric_ci('TranscriptFormer_mean_embedding_pseudobulk','donor_holdout')} for TranscriptFormer and {metric_ci('scVI_mean_embedding_pseudobulk','donor_holdout')} for scVI, indicating that donor-level aggregation changed prediction behavior.")
    p(doc, f"Dataset-holdout AUROC was lower and more variable: {metric_ci('metadata_confounder_logistic','dataset_holdout')} for metadata alone, {metric_ci('TranscriptFormer_embedding_logistic','dataset_holdout')} for TranscriptFormer, and {metric_ci('scVI_embedding_logistic','dataset_holdout')} for scVI. These results indicate that apparent aging-related prediction was sensitive to study composition and dataset shift.")
    pic(doc, "outputs/upgrade_figures/figure5_30seed_leakage_diagnostics.png", "Figure 6. Leakage and label-structure diagnostics. Generalization drop is defined as random-cell AUROC minus donor-holdout AUROC. Permutation gap is defined as true-label random-cell AUROC minus shuffled-label random-cell AUROC under the same split structure. These diagnostics assess whether apparent prediction persists after donor-aware evaluation and exceeds label-shuffled controls.", width=6.3)

    doc.add_heading("Question 4: how competitive are metadata-only baselines and what confounding structure is visible?", level=2)
    p(doc, "Age labels were unevenly distributed across dataset, disease, assay, and sex. The metadata-only coefficient analysis identified source and assay categories with large old-versus-young coefficients, and embedding principal components also showed visible alignment with non-age metadata. These diagnostics support treating metadata baselines as an essential benchmark component.")
    pic(doc, "outputs/upgrade_figures/figure6_age_label_confounding_structure.png", "Figure 7. Metadata structure of the age labels. Donor-level young and old labels are summarized across dataset, disease, assay, and sex categories in the blood T-cell benchmark sample. Uneven category distributions indicate potential metadata-derived predictability and motivate source-aware baselines.", width=6.3)
    pic(doc, "outputs/upgrade_figures/figure6b_metadata_coefficients.png", "Figure 8. Largest metadata-only coefficients for old-versus-young classification.")

    doc.add_heading("Question 5: what do disease-free and LODO diagnostics show about confounding and external generalization?", level=2)
    if not lodo.empty:
        best = lodo.sort_values("auroc", ascending=False).iloc[0]
        p(doc, f"Only one dataset met LODO eligibility after requiring both age classes in the held-out set. This held-out dataset contained {int(best['heldout_donors'])} donors and {int(best['heldout_cells'])} cells. Across the eligible LODO test, AUROC ranged from {lodo['auroc'].min():.3f} to {lodo['auroc'].max():.3f} across models, with the highest value observed for {best['model']}. The remaining {len(lodo_excl)} datasets were excluded because they lacked sufficient young and old donors. LODO feasibility is therefore itself a diagnostic result, indicating that public metadata provided limited cross-study validation capacity for this stratum.")
    else:
        p(doc, "No dataset met LODO eligibility after requiring both age classes in the held-out set. All exclusions were recorded.")
    p(doc, f"In the disease-free donor-holdout subset, cell-level AUROC was {metric_ci('TranscriptFormer_embedding_logistic','disease_free_donor_holdout')} for TranscriptFormer and {metric_ci('scVI_embedding_logistic','disease_free_donor_holdout')} for scVI. Donor-level mean embeddings increased performance to {metric_ci('TranscriptFormer_mean_embedding_pseudobulk','disease_free_donor_holdout')} and {metric_ci('scVI_mean_embedding_pseudobulk','disease_free_donor_holdout')}, respectively. This result describes benchmark behavior after disease-status restriction within the available public metadata and does not establish a biological aging mechanism.")
    p(doc, "High disease-free donor-holdout AUROC does not remove residual dataset, assay, recruitment, or cohort effects. Disease-free filtering reduces one observed source of confounding, but donors can remain structured by study origin, technical protocol, source population, and metadata completeness.")
    pic(doc, "outputs/upgrade_figures/figure7_lodo_dataset_holdout.png", "Figure 9. Leave-one-dataset-out diagnostic for the eligible held-out dataset. Each bar shows model performance when the only dataset satisfying the predefined young-old balance rule was held out for testing. The small number of eligible held-out datasets is part of the diagnostic finding and reflects limited cross-study validation capacity in the public metadata.", width=6.3)

    doc.add_heading("Question 6: how sensitive are donor-mean embeddings to sampling depth?", level=2)
    p(doc, f"Sampling-depth sensitivity showed monotonic gains from 1 to 8 cells per donor. TranscriptFormer donor-mean AUROC under donor holdout increased from {depth_metric('TranscriptFormer_mean_embedding_pseudobulk','donor_holdout',1):.3f} at one cell per donor to {depth_metric('TranscriptFormer_mean_embedding_pseudobulk','donor_holdout',8):.3f} at eight cells per donor. The requested 16-cell point was identical to the eight-cell point because the benchmark sample was capped at eight cells per donor.")
    pic(doc, "outputs/upgrade_figures/figure8_sampling_depth_sensitivity.png", "Figure 10. Sampling-depth sensitivity of donor-mean embeddings. Donor-mean scVI and TranscriptFormer embeddings were evaluated after sampling 1, 2, 4, 8, or 16 cells per donor. The 16-cell condition is capped by the existing eight-cell-per-donor benchmark sample, so it should be read as a capped sensitivity point rather than a new deeper-sampling experiment.", width=6.3)

    doc.add_paragraph("Table 1. Selected 30-seed AUROC results.")
    selected = bench[
        bench["model"].isin([
            "metadata_confounder_logistic",
            "scVI_embedding_logistic",
            "TranscriptFormer_embedding_logistic",
            "scVI_mean_embedding_pseudobulk",
            "TranscriptFormer_mean_embedding_pseudobulk",
        ])
        & bench["split"].isin(["random_cell", "donor_holdout", "dataset_holdout", "disease_free_donor_holdout"])
    ].copy()
    labels = {
        "metadata_confounder_logistic": "Metadata",
        "scVI_embedding_logistic": "scVI cell",
        "TranscriptFormer_embedding_logistic": "TF cell",
        "scVI_mean_embedding_pseudobulk": "scVI donor mean",
        "TranscriptFormer_mean_embedding_pseudobulk": "TF donor mean",
    }
    selected["Model"] = selected["model"].map(labels)
    selected["Split"] = selected["split"].str.replace("_", " ", regex=False)
    selected["AUROC mean"] = selected["auroc_mean"].round(3)
    selected["95% CI"] = selected.apply(lambda r: f"{r['auroc_ci95_low']:.3f}-{r['auroc_ci95_high']:.3f}", axis=1)
    selected["AUPRC mean"] = selected["auprc_mean"].round(3)
    table(doc, selected[["Model", "Split", "AUROC mean", "95% CI", "AUPRC mean"]], ["Model", "Split", "AUROC mean", "95% CI", "AUPRC mean"])

    doc.add_heading("Discussion", level=1)
    p(doc, "This manuscript presents a leakage-aware computational benchmark for public single-cell atlas reuse. Its main contribution is to show that prediction behavior changes substantially when evaluation moves from random-cell splitting to donor-, disease-, and study-aware diagnostics. This practical point is important for computational biology studies that reuse public CELLxGENE resources for supervised prediction.")
    p(doc, "Random-cell performance alone was insufficient. Metadata-only prediction reached a random-cell AUROC near 0.77, exceeding TranscriptFormer and scVI cell-level models under the same split. This indicates that source variables carried substantial age-related information in the benchmark sample. The pattern is consistent with prior single-cell integration work showing that batch, assay, and dataset structure can remain visible after representation learning or correction [27-29]. Related atlas-integration benchmarks also emphasize that dataset composition can influence downstream comparisons [30,31]. For public biological prediction benchmarks, metadata-only baselines should therefore be treated as standard controls rather than optional sensitivity analyses.")
    p(doc, "Dataset-aware evaluation provided the strongest dataset-shift diagnostic. Cell-level scVI and TranscriptFormer performance was lower under grouped dataset holdout, and the metadata baseline was also unstable. Public atlases aggregate studies with different recruitment frames, protocols, and disease composition, so donor-holdout evaluation alone cannot establish study-level transfer [14-16]. CELLxGENE-hosted resources make this evaluation practical, but do not remove source imbalance by themselves [17]. This interpretation is consistent with biomedical machine-learning studies that emphasize leakage control and external validation [39-41]. Transparent prediction-model reporting guidance makes the same point for validation design [42]. It also aligns with dataset-shift work showing that performance can change when deployment data differ from development data [46,47].")
    p(doc, "LODO feasibility was itself an informative result. Only one dataset contained enough young and old donors to support a held-out evaluation, while most datasets were excluded for age-class imbalance. This means that nominal dataset counts can overstate the effective external-validation structure available in public metadata. Reporting LODO exclusions therefore helps readers distinguish a large public atlas from a well-balanced cross-study benchmark.")
    p(doc, "Disease-free donor-holdout analyses showed strong performance for donor-mean embeddings, but the disease-free subset remains a public-metadata-derived cohort. Removing donors with overt disease annotations reduces one confounding axis, while residual dataset, assay, recruitment, cohort, and metadata-completeness effects can remain. The disease-free result therefore supports the need for layered diagnostics rather than replacing dataset-aware validation.")
    p(doc, "Donor-mean aggregation improved benchmark stability, particularly in disease-free donor-holdout analyses and sampling-depth sensitivity. Donor-level averaging may reduce cell-level sampling noise and better match the donor-level age label. The study boundary is that this is a computational observation about label unit, sampling depth, and hosted representation behavior, not evidence of a biological aging mechanism or a clinically validated predictor.")
    p(doc, "The findings also clarify how hosted embeddings should be evaluated. scVI and TranscriptFormer provided useful representations for a public-data benchmark, but their apparent performance depended on split design, metadata structure, disease restriction, and aggregation level. Related models such as Geneformer, scGPT, and scFoundation were discussed as context, but were not directly benchmarked here [20-22]. GeneCompass was also background context rather than a directly evaluated model [23]. The manuscript should therefore be read as a hosted representation assessment for scVI and TranscriptFormer in blood T cells, not as a comparison of all single-cell foundation models.")
    p(doc, "Several limitations remain. The completed empirical benchmark is restricted to blood T cells. LODO evaluation was feasible for only one eligible held-out dataset. The study evaluates hosted scVI and TranscriptFormer embeddings, not all single-cell foundation models. No raw-expression benchmark or manual external metadata harmonization was added. Age was inferred from public development_stage labels rather than from a harmonized donor-age field. Repository DOI and archival links should be finalized before submission.")
    p(doc, "Despite these limitations, the study provides a reproducible framework for metadata-aware validation of public single-cell prediction benchmarks. Its contribution is a workflow that connects metadata feasibility auditing, leakage diagnostics, source-aware splitting, label-shuffling controls, and sampling-depth sensitivity. This framing is aligned with reporting and reproducibility guidance for biomedical prediction studies, where transparent validation design is as important as point-estimate performance [41,42,49].")

    doc.add_heading("Conclusions", level=1)
    p(doc, "This blood T-cell benchmark shows that hosted single-cell representations can support aging-related prediction in public CELLxGENE data, but only when interpreted through metadata-aware validation. Metadata-only prediction, dataset-holdout fragility, and LODO exclusions show that public single-cell benchmarks are strongly constrained by donor, disease, and study structure. The proposed workflow provides a reproducible framework for leakage-aware hosted representation assessment rather than a pan-tissue human aging predictor.")

    doc.add_heading("Supplementary Information", level=1)
    p(doc, "Supplementary tables include the metadata feasibility map, threshold sensitivity analysis, 30-seed benchmark long-format results, summary metrics, paired comparisons, leakage diagnostics, LODO exclusions, confounding source data, sampling-depth sensitivity, extra-stratum feasibility audit, and upgrade audit JSON.")
    doc.add_paragraph("Supplementary Table S1. Benchmark risks and diagnostic controls.")
    risk_table = pd.DataFrame(
        [
            {"Risk": "Donor leakage", "Diagnostic control": "Donor-holdout split", "Interpretation": "Tests whether performance persists when donors are not shared across train and test."},
            {"Risk": "Disease confounding", "Diagnostic control": "Disease-free analysis and disease metadata diagnostics", "Interpretation": "Tests whether prediction behavior changes after disease-status restriction."},
            {"Risk": "Study or dataset shift", "Diagnostic control": "Grouped dataset-holdout and LODO feasibility", "Interpretation": "Tests whether prediction behavior transfers across dataset_id groups."},
            {"Risk": "Metadata-only predictability", "Diagnostic control": "Metadata-only baseline", "Interpretation": "Quantifies age-related prediction from dataset, assay, disease, and sex alone."},
            {"Risk": "Spurious label structure", "Diagnostic control": "Label-shuffling control", "Interpretation": "Checks whether prediction exceeds shuffled-label baselines under the same split structure."},
            {"Risk": "Donor aggregation sensitivity", "Diagnostic control": "Sampling-depth sensitivity", "Interpretation": "Tests how donor-mean embeddings change as cells per donor vary."},
        ]
    )
    table(doc, risk_table, ["Risk", "Diagnostic control", "Interpretation"])
    doc.add_paragraph("Supplementary Table S2. LODO eligibility and exclusion summary.")
    eligible_dataset_n = int(lodo["dataset_id"].nunique()) if not lodo.empty else 0
    excluded_dataset_n = int(len(lodo_excl))
    no_young = int((lodo_excl["young_donors"] < 2).sum()) if not lodo_excl.empty else 0
    no_old = int((lodo_excl["old_donors"] < 2).sum()) if not lodo_excl.empty else 0
    train_test_issue = int(lodo_excl["exclusion_reason"].astype(str).str.contains("train or test set lacks both age classes").sum()) if not lodo_excl.empty else 0
    lodo_summary = pd.DataFrame(
        [
            {"Category": "Eligible held-out datasets", "Dataset count": eligible_dataset_n, "Interpretation": "Had at least two young and two old donors and valid remaining training classes."},
            {"Category": "Excluded held-out datasets", "Dataset count": excluded_dataset_n, "Interpretation": "Failed LODO eligibility under the predefined age-class balance rule."},
            {"Category": "Excluded with fewer than two young donors", "Dataset count": no_young, "Interpretation": "Held-out dataset lacked sufficient young donors."},
            {"Category": "Excluded with fewer than two old donors", "Dataset count": no_old, "Interpretation": "Held-out dataset lacked sufficient old donors."},
            {"Category": "Excluded with invalid train/test class balance", "Dataset count": train_test_issue, "Interpretation": "Held-out or remaining training set lacked both age classes."},
        ]
    )
    table(doc, lodo_summary, ["Category", "Dataset count", "Interpretation"])
    p(doc, "The complete dataset-level LODO table with dataset identifiers, donor counts, cell counts, and exclusion reasons is provided as `analysis/upgrade_outputs/lodo_dataset_exclusions.tsv` and `analysis/upgrade_outputs/lodo_dataset_holdout_metrics.tsv`.")
    p(doc, "Supplementary Checklist S1 maps each diagnostic to a leakage or confounding risk: donor-holdout for donor leakage; disease-free analysis for disease confounding; grouped dataset-holdout and LODO for study shift; metadata-only baseline for metadata predictability; label shuffling for spurious label structure; and sampling-depth sensitivity for donor aggregation.")
    doc.add_heading("Acknowledgements", level=1)
    p(doc, "[UNRESOLVED PLACEHOLDER: Acknowledgements should be completed before submission.]")
    doc.add_heading("Authors' contributions", level=1)
    p(doc, "[UNRESOLVED PLACEHOLDER: Author contributions should be completed before submission.]")
    doc.add_heading("Funding", level=1)
    p(doc, "[UNRESOLVED PLACEHOLDER: Funding information should be completed before submission.]")
    doc.add_heading("Availability of data and materials", level=1)
    p(doc, "The analyses used CELLxGENE Census release 2025-11-08 and hosted embeddings from the same release. Project scripts, split definitions or deterministic split-generation code, target-source tables, QC outputs, benchmark outputs, figure source data, candidate-level feasibility results, and the output manifest are included in the reproducibility package and supplementary materials. The public GitHub repository is available at https://github.com/seefreewind/cellxgene-blood-tcell-aging-benchmark. [UNRESOLVED PLACEHOLDER: Zenodo DOI to be inserted before submission.] Public archival availability should not be inferred until the Zenodo DOI is finalized.")
    doc.add_heading("Ethics approval and consent to participate", level=1)
    p(doc, "This study used public de-identified single-cell metadata and hosted embeddings. No new human participants or animal experiments were included.")
    doc.add_heading("Consent for publication", level=1)
    p(doc, "Not applicable.")
    doc.add_heading("Competing interests", level=1)
    p(doc, "[UNRESOLVED PLACEHOLDER: Competing interests declaration should be completed before submission.]")

    doc.add_heading("References", level=1)
    refs = [
        "Horvath S. DNA methylation age of human tissues and cell types. Genome Biology. 2013;14:R115. doi:10.1186/gb-2013-14-10-r115.",
        "Hannum G, Guinney J, Zhao L, Zhang L, Hughes G, Sadda S, et al. Genome-wide methylation profiles reveal quantitative views of human aging rates. Molecular Cell. 2013;49:359-367. doi:10.1016/j.molcel.2012.10.016.",
        "Levine ME, Lu AT, Quach A, Chen BH, Assimes TL, Bandinelli S, et al. An epigenetic biomarker of aging for lifespan and healthspan. Aging. 2018;10:573-591. doi:10.18632/aging.101414.",
        "Lu AT, Quach A, Wilson JG, Reiner AP, Aviv A, Raj K, et al. DNA methylation GrimAge strongly predicts lifespan and healthspan. Aging. 2019;11:303-327. doi:10.18632/aging.101684.",
        "Craig T, Smelick C, Tacutu R, Wuttke D, Wood SH, Stanley H, et al. The Digital Ageing Atlas: integrating the diversity of age-related changes into a unified resource. Nucleic Acids Research. 2015;43:D873-D878. doi:10.1093/nar/gku843.",
        "Aging Atlas Consortium. Aging Atlas: a multi-omics database for aging biology. Nucleic Acids Research. 2021;49:D825-D830. doi:10.1093/nar/gkaa894.",
        "Tabula Muris Consortium. A single-cell transcriptomic atlas characterizes ageing tissues in the mouse. Nature. 2020;583:590-595. doi:10.1038/s41586-020-2496-1.",
        "Zou Z, Long X, Zhao Q, Zheng Y, Song M, Ma S, et al. A single-cell transcriptomic atlas of human skin aging. Developmental Cell. 2021;56:383-397.e8. doi:10.1016/j.devcel.2020.11.002.",
        "Lu Y, Brommer B, Tian X, Krishnan A, Meer M, Wang C, et al. Reprogramming to recover youthful epigenetic information and restore vision. Nature. 2020;588:124-129. doi:10.1038/s41586-020-2975-4.",
        "Zakar-Polyak E, Csordas A, Palovics R, Kerepesi C. Profiling the transcriptomic age of single-cells in humans. Communications Biology. 2024;7:1397. doi:10.1038/s42003-024-07094-5.",
        "Kedlian VR, Wang Y, Herder C, et al. Human skeletal muscle aging atlas. Nature Aging. 2024;4:727-744. doi:10.1038/s43587-024-00613-3.",
        "Bartz J, Ma X, Zhang L, Dong X. Human Cell Aging Transcriptome Atlas (HCATA): a single-cell atlas of age-associated transcriptomic alterations across human tissues. Communications Biology. 2025. doi:10.1038/s42003-025-08845-8.",
        "Aging Fly Cell Atlas Consortium. Aging Fly Cell Atlas identifies exhaustive aging features at cellular resolution. Science. 2023;380:eadg0934. doi:10.1126/science.adg0934.",
        "Regev A, Teichmann SA, Lander ES, Amit I, Benoist C, Birney E, et al. The Human Cell Atlas. eLife. 2017;6:e27041. doi:10.7554/eLife.27041.",
        "Han X, Zhou Z, Fei L, Sun H, Wang R, Chen Y, et al. Construction of a human cell landscape at single-cell level. Nature. 2020;581:303-309. doi:10.1038/s41586-020-2157-4.",
        "Tabula Sapiens Consortium. The Tabula Sapiens: a multiple-organ, single-cell transcriptomic atlas of humans. Science. 2022;376:eabl4896. doi:10.1126/science.abl4896.",
        "Abdulla S, Aevermann BD, Assis P, Badajoz S, Bell SM, Bezzi E, et al. CZ CELLxGENE Discover: a single-cell data platform for scalable exploration, analysis and modeling of aggregated data. Nucleic Acids Research. 2025;53:D886-D900. doi:10.1093/nar/gkae1142.",
        "Lopez R, Regier J, Cole MB, Jordan MI, Yosef N. Deep generative modeling for single-cell transcriptomics. Nature Methods. 2018;15:1053-1058. doi:10.1038/s41592-018-0229-2.",
        "Gayoso A, Lopez R, Xing G, Boyeau P, Valiollah Pour Amiri V, Hong J, et al. A Python library for probabilistic analysis of single-cell omics data. Nature Biotechnology. 2022;40:163-166. doi:10.1038/s41587-021-01206-w.",
        "Theodoris CV, Xiao L, Chopra A, Chaffin MD, Al Sayed ZR, Hill MC, et al. Transfer learning enables predictions in network biology. Nature. 2023;618:616-624. doi:10.1038/s41586-023-06139-9.",
        "Cui H, Wang C, Maan H, Pang K, Luo F, Duan N, et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods. 2024;21:1470-1480. doi:10.1038/s41592-024-02201-0.",
        "Hao M, Gong J, Zeng X, Liu C, Guo Y, Cheng X, et al. Large-scale foundation model on single-cell transcriptomics. Nature Methods. 2024;21:1481-1491. doi:10.1038/s41592-024-02305-7.",
        "Yang X, Liu G, Feng G, et al. GeneCompass: deciphering universal gene regulatory mechanisms with a knowledge-informed cross-species foundation model. Cell Research. 2024;34:830-845. doi:10.1038/s41422-024-01034-y.",
        "Pearce JD, Simmonds SE, Mahmoudabadi G, Krishnan L, Palla G, Istrate AM, et al. TranscriptFormer: a generative cell atlas across 1.5 billion years of evolution. Science. 2026;eaec8514. doi:10.1126/science.aec8514.",
        "Luecken MD, Theis FJ. Current best practices in single-cell RNA-seq analysis: a tutorial. Molecular Systems Biology. 2019;15:e8746. doi:10.15252/msb.20188746.",
        "Stuart T, Butler A, Hoffman P, Hafemeister C, Papalexi E, Mauck WM 3rd, et al. Comprehensive integration of single-cell data. Cell. 2019;177:1888-1902.e21. doi:10.1016/j.cell.2019.05.031.",
        "Haghverdi L, Lun ATL, Morgan MD, Marioni JC. Batch effects in single-cell RNA-sequencing data are corrected by matching mutual nearest neighbors. Nature Biotechnology. 2018;36:421-427. doi:10.1038/nbt.4091.",
        "Korsunsky I, Millard N, Fan J, Slowikowski K, Zhang F, Wei K, et al. Fast, sensitive and accurate integration of single-cell data with Harmony. Nature Methods. 2019;16:1289-1296. doi:10.1038/s41592-019-0619-0.",
        "Johnson WE, Li C, Rabinovic A. Adjusting batch effects in microarray expression data using empirical Bayes methods. Biostatistics. 2007;8:118-127. doi:10.1093/biostatistics/kxj037.",
        "Tran HTN, Ang KS, Chevrier M, Zhang X, Lee NYS, Goh M, et al. A benchmark of batch-effect correction methods for single-cell RNA sequencing data. Genome Biology. 2020;21:12. doi:10.1186/s13059-019-1850-9.",
        "Luecken MD, Buttner M, Chaichoompu K, Danese A, Interlandi M, Mueller MF, et al. Benchmarking atlas-level data integration in single-cell genomics. Nature Methods. 2022;19:41-50. doi:10.1038/s41592-021-01336-8.",
        "Polanski K, Young MD, Miao Z, Meyer KB, Teichmann SA, Park JE. BBKNN: fast batch alignment of single cell transcriptomes. Bioinformatics. 2020;36:964-965. doi:10.1093/bioinformatics/btz625.",
        "Hie B, Bryson B, Berger B. Efficient integration of heterogeneous single-cell transcriptomes using Scanorama. Nature Biotechnology. 2019;37:685-691. doi:10.1038/s41587-019-0113-3.",
        "Buttner M, Miao Z, Wolf FA, Teichmann SA, Theis FJ. A test metric for assessing single-cell RNA-seq batch correction. Nature Methods. 2019;16:43-49. doi:10.1038/s41592-018-0254-1.",
        "Amezquita RA, Lun ATL, Becht E, Carey VJ, Carpp LN, Geistlinger L, et al. Orchestrating single-cell analysis with Bioconductor. Nature Methods. 2020;17:137-145. doi:10.1038/s41592-019-0654-x.",
        "Wolf FA, Angerer P, Theis FJ. SCANPY: large-scale single-cell gene expression data analysis. Genome Biology. 2018;19:15. doi:10.1186/s13059-017-1382-0.",
        "Virshup I, Rybakov S, Theis FJ, Angerer P, Wolf FA. anndata: annotated data. bioRxiv. 2021. doi:10.1101/2021.12.16.473007.",
        "Pedregosa F, Varoquaux G, Gramfort A, Michel V, Thirion B, Grisel O, et al. Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research. 2011;12:2825-2830.",
        "Varoquaux G. Cross-validation failure: small sample sizes lead to large error bars. NeuroImage. 2018;180:68-77. doi:10.1016/j.neuroimage.2017.06.061.",
        "Vabalas A, Gowen E, Poliakoff E, Casson AJ. Machine learning algorithm validation with a limited sample size. PLoS ONE. 2019;14:e0224365. doi:10.1371/journal.pone.0224365.",
        "Collins GS, Reitsma JB, Altman DG, Moons KGM. Transparent reporting of a multivariable prediction model for individual prognosis or diagnosis (TRIPOD): the TRIPOD statement. Annals of Internal Medicine. 2015;162:55-63. doi:10.7326/M14-0697.",
        "Luo W, Phung D, Tran T, Gupta S, Rana S, Karmakar C, et al. Guidelines for developing and reporting machine learning predictive models in biomedical research: a multidisciplinary view. Journal of Medical Internet Research. 2016;18:e323. doi:10.2196/jmir.5870.",
        "Kelly CJ, Karthikesalingam A, Suleyman M, Corrado G, King D. Key challenges for delivering clinical impact with artificial intelligence. BMC Medicine. 2019;17:195. doi:10.1186/s12916-019-1426-2.",
        "Roberts M, Driggs D, Thorpe M, Gilbey J, Yeung M, Ursprung S, et al. Common pitfalls and recommendations for using machine learning to detect and prognosticate for COVID-19 using chest radiographs and CT scans. Nature Machine Intelligence. 2021;3:199-217. doi:10.1038/s42256-021-00307-0.",
        "Wynants L, Van Calster B, Collins GS, Riley RD, Heinze G, Schuit E, et al. Prediction models for diagnosis and prognosis of covid-19 infection: systematic review and critical appraisal. BMJ. 2020;369:m1328. doi:10.1136/bmj.m1328.",
        "Koh PW, Sagawa S, Marklund H, Xie SM, Zhang M, Balsubramani A, et al. WILDS: a benchmark of in-the-wild distribution shifts. Proceedings of Machine Learning Research. 2021;139:5637-5664.",
        "Bzdok D, Engemann DA, Thirion B. Inference and prediction diverge in biomedicine. Patterns. 2020;1:100119. doi:10.1016/j.patter.2020.100119.",
        "Riley RD, Ensor J, Snell KIE, Harrell FE Jr, Martin GP, Reitsma JB, et al. Calculating the sample size required for developing a clinical prediction model. BMJ. 2020;368:m441. doi:10.1136/bmj.m441.",
        "Open Science Collaboration. Estimating the reproducibility of psychological science. Science. 2015;349:aac4716. doi:10.1126/science.aac4716.",
    ]
    for item in refs:
        ref(doc, item)

    for section in doc.sections:
        add_line_numbers(section)
    OUT.mkdir(parents=True, exist_ok=True)
    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
