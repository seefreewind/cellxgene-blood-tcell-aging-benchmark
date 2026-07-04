from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT = ROOT / "outputs"
DOCX = OUT / "donor_disease_aware_sc_aging_benchmark_draft.docx"
MD = ROOT / "manuscript" / "donor_disease_aware_sc_aging_benchmark_draft.md"
OUT.mkdir(exist_ok=True)

system = pd.read_csv(ROOT / "analysis/outputs/cell_system_feasibility_map.tsv", sep="\t")
tissue = pd.read_csv(ROOT / "analysis/outputs/metadata_feasibility_map.tsv", sep="\t")
summary = pd.read_json(ROOT / "analysis/outputs/metadata_audit_summary.json", typ="series")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_from_df(doc: Document, df: pd.DataFrame, columns: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        set_cell_text(hdr[i], col, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.3f}" if value < 1 else f"{value:,.0f}"
            elif isinstance(value, int):
                value = f"{value:,}"
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()


def add_reference(doc: Document, n: int, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def build_markdown() -> str:
    top = tissue.head(12)
    system_lines = []
    for _, r in system.sort_values("donor_n", ascending=False).iterrows():
        system_lines.append(
            f"- {r['cell_type']}: {int(r['donor_n']):,} donors, "
            f"{int(r['cell_n']):,} cells, {int(r['tissue_n']):,} tissues, "
            f"{int(r['dataset_n']):,} datasets, disease-free donor fraction "
            f"{r['disease_free_donor_fraction']:.3f}, status {r['feasibility_status']}."
        )
    tissue_lines = []
    for _, r in top.iterrows():
        tissue_lines.append(
            f"- {r['tissue_general']} / {r['cell_type']}: {int(r['donor_n']):,} donors, "
            f"{int(r['cell_n']):,} cells, {int(r['dataset_n']):,} datasets, "
            f"disease-free donor fraction {r['disease_free_donor_fraction']:.3f}."
        )
    return f"""# Donor- and disease-aware benchmarking of single-cell foundation models for human cellular aging

## Draft status

This is a first Word manuscript draft generated from the current project plan and a Phase 0 CELLxGENE Census metadata audit. Model benchmarking results have not yet been generated; the Results section therefore separates completed metadata-audit findings from planned benchmark outputs.

## One-sentence argument

In public human single-cell aging resources, we show that donor-, disease-, and study-aware metadata auditing is a necessary first step before foundation-model benchmarking, supported by a CELLxGENE Census audit of five target cell systems and bounded by the absence of completed model-performance and external-validation analyses in this first draft.

## Phase 0 audit summary

- Census version: {summary['census_version']}
- Target primary cell metadata rows: {int(summary['target_primary_cell_rows']):,}
- Unique donors: {int(summary['unique_donors']):,}
- Unique datasets: {int(summary['unique_datasets']):,}
- Unique collections: {int(summary['unique_collections']):,}
- Unique tissues: {int(summary['unique_tissues']):,}
- Unique target cell types: {int(summary['unique_cell_types']):,}

### Cell-system feasibility

{chr(10).join(system_lines)}

### Top eligible tissue-cell strata

{chr(10).join(tissue_lines)}
"""


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Donor- and disease-aware benchmarking of single-cell foundation models for human cellular aging")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("First manuscript draft with completed Phase 0 metadata feasibility audit").italic = True

    doc.add_paragraph("Author list: [To be completed]")
    doc.add_paragraph("Affiliations: [To be completed]")
    doc.add_paragraph("Corresponding author: [To be completed]")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("Background: Public single-cell resources are increasingly used to study human aging, and foundation models are now being applied to predict cell states from transcriptomic profiles. However, apparent cellular aging prediction may be inflated by donor overlap, disease imbalance, tissue composition, study origin, and assay effects. A benchmark that starts from metadata feasibility rather than assuming pan-tissue comparability is therefore required.")
    doc.add_paragraph(f"Results: We first performed a metadata feasibility audit using CELLxGENE Census version {summary['census_version']}. Across five target cell systems, the audit captured {int(summary['target_primary_cell_rows']):,} primary-cell metadata rows from {int(summary['unique_donors']):,} donors, {int(summary['unique_datasets']):,} datasets, {int(summary['unique_collections']):,} collections, and {int(summary['unique_tissues']):,} tissues. At the system level, T cells, macrophages, and fibroblasts met the predefined eligible-stratum criteria, whereas endothelial cells and monocytes were classified as pilot candidates because disease-free donor retention was below the 50% threshold. Several tissue-cell strata, including blood T cells, breast macrophages, breast fibroblasts, eye fibroblasts, and liver macrophages, also met donor, age, dataset, and disease-free criteria.")
    doc.add_paragraph("Conclusions: These findings support a staged benchmark focused first on immune, endothelial, and stromal systems rather than an unconditional pan-tissue atlas. The planned benchmark will compare simple pseudobulk and cell-level baselines, conventional single-cell representations, hosted foundation-model embeddings, and biological signatures under random, donor-holdout, study-holdout, disease-free, and label-shuffling evaluations.")
    doc.add_paragraph("Keywords: single-cell RNA sequencing; cellular aging; foundation models; donor holdout; disease confounding; CELLxGENE Census; benchmark")

    doc.add_heading("Background", level=1)
    doc.add_paragraph("Single-cell RNA sequencing has made it possible to study aging at the level of individual cell types, tissues, and cellular states. Large public resources now aggregate millions of cells across studies, but the metadata structure of these resources is not neutral: older donors may be overrepresented in disease cohorts, specific tissues may be tied to particular studies, and assay or collection effects may correlate with age. These issues are especially important for cellular aging prediction, where a model can appear accurate if it learns donor identity, disease state, study origin, or tissue composition rather than transferable aging biology [1,2].")
    doc.add_paragraph("Single-cell foundation models have introduced a new layer of representation learning for transcriptomic data. Geneformer used transfer learning on large-scale single-cell transcriptomes to support predictions in network biology, and scGPT extended generative pretraining ideas to single-cell multi-omics representations [3,4]. These models are attractive for aging research because they may encode reusable cellular programs. At the same time, their scale and flexibility make leakage-aware evaluation essential. A foundation-model embedding that performs well under random cell splits may still fail when donors, studies, or disease states are held out.")
    doc.add_paragraph("The central premise of this study is therefore methodological and biological. Before asking whether a model predicts cellular aging, public metadata must show that the task can be evaluated with donor-, disease-, and study-aware controls. We propose a staged benchmark that starts with a feasibility audit, escalates only when metadata support the design, and interprets only those aging programs that remain stable under strict holdout and disease-free analyses.")

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design", level=2)
    doc.add_paragraph("This study is designed as a public-data benchmark and metadata audit of human cellular aging prediction. The primary resource is CELLxGENE Census because it provides programmatic access to standardized single-cell metadata and expression matrices. HCATA will be used as an external aging-oriented reference for biological interpretation, Aging Atlas will be used for aging-related gene set context, and SenNet will be treated as an optional senescence-annotation resource where dataset access and metadata completeness allow [1,2,5,6].")
    doc.add_heading("Phase 0 metadata feasibility audit", level=2)
    doc.add_paragraph(f"The Phase 0 audit used CELLxGENE Census release {summary['census_version']}. The audit queried primary human-cell metadata for five target cell systems: T cells, monocytes, macrophages, endothelial cells, and fibroblasts. Retrieved fields included dataset identifier, collection identifier, assay, tissue, tissue_general, cell_type, disease, sex, development_stage, donor_id, suspension_type, and is_primary_data. Dataset-level collection metadata were joined from the Census dataset table.")
    doc.add_paragraph("Age labels were classified from development_stage into young_or_adult, old_or_elderly, annotated_other, or unknown using conservative text matching. Disease-free donor retention was estimated from donors annotated as normal or healthy. These labels are proxies and must be replaced by precise numeric donor age fields where source datasets provide them.")
    doc.add_paragraph("A tissue-cell stratum was considered eligible when it contained at least 30 donors, at least 10 young/adult donors, at least 10 old/elderly donors, at least two datasets, and at least 50% disease-free donor retention. Strata with at least 15 donors, at least two datasets, and at least 30% disease-free donor retention were retained as pilot candidates. These thresholds were chosen to prevent model evaluation from being driven by single-dataset, single-disease, or single-age-composition artifacts.")
    doc.add_heading("Planned benchmark models", level=2)
    doc.add_paragraph("The benchmark will be organized into three model layers. The first layer will include simple baselines: pseudobulk ridge regression, pseudobulk elastic net, cell-level logistic regression, PCA plus linear models, and highly variable gene mean-expression baselines. The second layer will include conventional single-cell representations, including PCA, scVI, and batch-corrected embeddings as sensitivity analyses. The third layer will use hosted or released foundation-model embeddings, prioritizing CELLxGENE-hosted embeddings before any project-specific large-scale inference.")
    doc.add_heading("Planned evaluation splits", level=2)
    doc.add_paragraph("Five split designs will be used. Random cell splits will estimate apparent performance. Donor-holdout splits will prevent cells from the same donor appearing in training and test data. Study- or dataset-holdout splits will evaluate transfer across studies, laboratories, and processing protocols. Disease-free or disease-stratified splits will test whether age prediction persists after disease-associated expression programs are controlled. Label-shuffling analyses will preserve donor, tissue, and dataset structure while permuting age labels to estimate the performance expected from structured leakage alone.")
    doc.add_heading("Planned outcomes", level=2)
    doc.add_paragraph("For continuous age prediction, the primary metrics will be mean absolute error and Spearman correlation. For young-old classification, the primary metrics will be AUROC, AUPRC, and balanced accuracy. Two derived quantities will be emphasized: generalization drop, defined as random-split performance minus holdout performance, and permutation gap, defined as true-label performance minus shuffled-label performance.")

    doc.add_heading("Results", level=1)
    doc.add_heading("Metadata audit supports a staged rather than unconditional pan-tissue benchmark", level=2)
    doc.add_paragraph(f"The Phase 0 audit retrieved {int(summary['target_primary_cell_rows']):,} primary-cell metadata rows across five target cell systems. These rows represented {int(summary['unique_donors']):,} donors, {int(summary['unique_datasets']):,} datasets, {int(summary['unique_collections']):,} collections, and {int(summary['unique_tissues']):,} tissue_general labels. This scale supports benchmark development, but the distribution across cell systems and disease-free subsets was uneven.")
    doc.add_paragraph("At the cell-system level, T cells included 3,791 donors and 1,267,019 cells across 56 tissues and 119 datasets, with a disease-free donor fraction of 0.548. Macrophages included 2,984 donors and 1,736,309 cells across 66 tissues and 125 datasets, with disease-free donor retention of 0.504. Fibroblasts included 2,214 donors and 2,663,513 cells across 61 tissues and 212 datasets, with disease-free donor retention of 0.586. These three systems met the predefined eligible-stratum criteria. Endothelial cells and monocytes had broad donor and dataset coverage but disease-free donor fractions below 0.50, so they were retained as pilot candidates rather than primary benchmark strata.")

    fig1 = ROOT / "outputs/figures/figure1_cell_system_feasibility.png"
    if fig1.exists():
        doc.add_picture(str(fig1), width=Inches(6.2))
        cap = doc.add_paragraph("Figure 1. Cell-system-level feasibility audit for five target cell systems. Bars show donor coverage and disease-free donor retention; the dashed line marks the 50% disease-free retention threshold.")
        cap.runs[0].italic = True

    doc.add_heading("Eligible tissue-cell strata are concentrated in immune and stromal systems", level=2)
    doc.add_paragraph("Several tissue-cell combinations met the current eligibility criteria. The largest eligible stratum was blood T cells, with 824 donors, 125,501 cells, 20 datasets, and a disease-free donor fraction of 0.873. Other eligible strata included breast macrophages, breast fibroblasts, eye fibroblasts, breast T cells, liver macrophages, small-intestine macrophages, small-intestine fibroblasts, colon fibroblasts, and bone marrow macrophages. These results support a benchmark centered on immune and stromal systems, with endothelial analyses treated as a sensitivity or pilot layer until disease-free retention and age balance are improved.")

    fig2 = ROOT / "outputs/figures/figure2_tissue_cell_feasibility_heatmap.png"
    if fig2.exists():
        doc.add_picture(str(fig2), width=Inches(6.2))
        cap = doc.add_paragraph("Figure 2. Top eligible tissue-cell strata in the Phase 0 metadata audit. Values indicate unique donor counts for eligible strata.")
        cap.runs[0].italic = True

    doc.add_paragraph("Table 1 summarizes system-level feasibility results.")
    table_df = system.sort_values("donor_n", ascending=False).copy()
    table_df["disease_free_donor_fraction"] = table_df["disease_free_donor_fraction"].round(3)
    table_df["status"] = table_df["feasibility_status"].map(
        {"eligible_stratum": "eligible", "pilot_candidate": "pilot"}
    ).fillna(table_df["feasibility_status"])
    table_df = table_df.rename(
        columns={
            "cell_type": "Cell system",
            "donor_n": "Donors",
            "cell_n": "Cells",
            "tissue_n": "Tissues",
            "dataset_n": "Datasets",
            "disease_free_donor_fraction": "Disease-free",
            "status": "Status",
        }
    )
    add_table_from_df(
        doc,
        table_df,
        [
            "Cell system",
            "Donors",
            "Cells",
            "Tissues",
            "Datasets",
            "Disease-free",
            "Status",
        ],
    )

    doc.add_heading("Benchmark results pending", level=2)
    doc.add_paragraph("Model-performance results have not yet been generated. The next analytical stage will run the minimum pilot benchmark in T cells, macrophages or monocytes, fibroblasts, and endothelial cells where feasible. This first draft therefore does not report AUROC, AUPRC, MAE, Spearman correlation, generalization drop, permutation gap, or biological reproducibility scores.")

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph("The Phase 0 audit shows that public single-cell metadata can support a donor- and disease-aware benchmark for selected human cellular aging systems, but it does not justify an unconditional pan-tissue benchmark at this stage. T cells, macrophages, and fibroblasts showed sufficient donor, dataset, age-proxy, and disease-free coverage to enter the main benchmark. Endothelial cells and monocytes had strong overall coverage but weaker disease-free retention, making them suitable for pilot or sensitivity analyses rather than primary claims.")
    doc.add_paragraph("This finding is important because random split performance alone is unlikely to answer whether a model has learned aging biology. When cells from the same donor or study appear in both training and testing sets, models can exploit repeated donor structure, disease status, tissue composition, or assay effects. The planned benchmark therefore treats donor-holdout, study-holdout, disease-free evaluation, and label-shuffling controls as primary analyses rather than supplementary checks.")
    doc.add_paragraph("The study has several current limitations. First, the Phase 0 audit used development_stage as an age proxy because exact numeric donor age is not uniformly available through the queried Census metadata fields. Second, disease-free status was inferred from normal or healthy labels, which may not capture all disease annotations or donor-level clinical context. Third, no expression matrices, hosted embeddings, or model-performance results have yet been analyzed in this first draft. These limitations are not defects to hide; they define the next analytical steps and the boundary of the current manuscript.")
    doc.add_paragraph("If the planned pilot shows that donor- and study-holdout performance remains above permutation baselines, the manuscript can advance toward a benchmark of immune, endothelial, and stromal aging systems. If performance collapses under these controls, the paper will still make a useful contribution by showing that apparent cellular aging prediction in public single-cell data is strongly inflated by metadata structure and disease composition.")

    doc.add_heading("Conclusions", level=1)
    doc.add_paragraph("A donor- and disease-aware metadata audit should precede single-cell foundation-model benchmarking in human cellular aging. In the current CELLxGENE Census audit, T cells, macrophages, and fibroblasts met predefined feasibility criteria, while endothelial cells and monocytes were retained as pilot candidates. These results support a staged benchmark focused on immune and stromal systems, with strict donor, study, disease-free, and permutation controls required before any claim of transferable cellular aging biology.")

    doc.add_heading("Supplementary Information", level=1)
    doc.add_paragraph("Supplementary Table 1: CELLxGENE Census cell-system feasibility map. Supplementary Table 2: tissue-cell feasibility map. Supplementary Table 3: donor-level metadata summary. These files are currently available in the project analysis output directory.")
    doc.add_heading("Acknowledgements", level=1)
    doc.add_paragraph("[To be completed.]")
    doc.add_heading("Authors' contributions", level=1)
    doc.add_paragraph("[To be completed according to author roles.]")
    doc.add_heading("Funding", level=1)
    doc.add_paragraph("[To be completed.]")
    doc.add_heading("Availability of data and materials", level=1)
    doc.add_paragraph("The Phase 0 metadata audit used CELLxGENE Census release 2025-11-08. Audit scripts and generated summary tables are included in this project directory. External resources planned for validation include HCATA, Aging Atlas, and SenNet, subject to dataset-level access and metadata completeness.")
    doc.add_heading("Ethics approval and consent to participate", level=1)
    doc.add_paragraph("This study uses publicly available de-identified metadata and planned public single-cell datasets. No new human participants or animal experiments were included in the current analysis.")
    doc.add_heading("Consent for publication", level=1)
    doc.add_paragraph("Not applicable.")
    doc.add_heading("Competing interests", level=1)
    doc.add_paragraph("The authors declare that they have no competing interests. [Revise if needed.]")

    doc.add_heading("References", level=1)
    references = [
        "Hilton J, et al. CZ CELLxGENE Discover is an online analytical platform and the largest repository of standardized single-cell data. 2024. doi:10.7490/f1000research.1119921.1.",
        "CELLxGENE Census documentation. CZ CELLxGENE Discover Census and hosted embeddings documentation. Available from: https://chanzuckerberg.github.io/cellxgene-census/.",
        "Theodoris CV, Xiao L, Chopra A, Chaffin MD, Al Sayed ZR, Hill MC, et al. Transfer learning enables predictions in network biology. Nature. 2023;618:616-624. doi:10.1038/s41586-023-06139-9.",
        "Cui H, Wang C, Maan H, Pang K, Luo F, Duan N, et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods. 2024. doi:10.1038/s41592-024-02201-0.",
        "Bartz J, Ma X, Zhang L, Dong X. Human Cell Aging Transcriptome Atlas (HCATA): a single-cell atlas of age-associated transcriptomic alterations across human tissues. Communications Biology. 2025. doi:10.1038/s42003-025-08845-8.",
        "Aging Atlas: a multi-omics database for aging biology. Nucleic Acids Research. 2021. doi:10.1093/nar/gkaa894.",
        "SenNet Portal: Build, Optimization and Usage. 2026. doi:10.64898/2026.02.06.704469.",
    ]
    for i, ref in enumerate(references, start=1):
        add_reference(doc, i, ref)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Internal Draft Notes", level=1)
    doc.add_paragraph("This page should be removed before journal submission. It records missing items for the next draft.")
    for item in [
        "Replace age proxy with exact numeric donor ages where source metadata permit.",
        "Run the minimum pilot benchmark using pseudobulk ridge, PCA plus logistic regression, hosted embeddings, and label-shuffling.",
        "Add full software versions for all benchmark models.",
        "Expand literature review toward approximately 50 references after model results are available.",
        "Verify every final reference against PubMed, Crossref, or publisher pages.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(DOCX)
    MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
