from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT = ROOT / "outputs"
DOCX = OUT / "donor_disease_aware_sc_aging_completed_pilot_v1.docx"
MD = ROOT / "manuscript" / "donor_disease_aware_sc_aging_completed_pilot_v1.md"

meta_summary = pd.read_json(ROOT / "analysis/outputs/metadata_audit_summary.json", typ="series")
system = pd.read_csv(ROOT / "analysis/outputs/cell_system_feasibility_with_ci.tsv", sep="\t")
tissue = pd.read_csv(ROOT / "analysis/outputs/metadata_feasibility_map.tsv", sep="\t")
sens = pd.read_csv(ROOT / "analysis/outputs/threshold_sensitivity.tsv", sep="\t")
bench_audit = pd.read_json(ROOT / "analysis/benchmark_outputs/blood_tcell_embedding_benchmark_audit.json", typ="series")
bench = pd.read_csv(ROOT / "analysis/benchmark_outputs/blood_tcell_embedding_benchmark_results_summary.tsv", sep="\t")
diag = pd.read_csv(ROOT / "analysis/benchmark_outputs/blood_tcell_embedding_leakage_diagnostics.tsv", sep="\t")


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def style_doc(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        add_line_numbers(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = False
        style.paragraph_format.line_spacing = 1.15


def p(doc: Document, text: str):
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing = 2.0
    return para


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc: Document, df: pd.DataFrame, cols: list[str]):
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, col in enumerate(cols):
        set_cell(tbl.rows[0].cells[i], col, True)
        shade(tbl.rows[0].cells[i], "E8EEF5")
    for _, row in df.iterrows():
        cells = tbl.add_row().cells
        for i, col in enumerate(cols):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.3f}"
            set_cell(cells[i], str(value))
    doc.add_paragraph()


def pic(doc: Document, path: str, caption: str, width: float = 6.2):
    fp = ROOT / path
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.paragraph_format.line_spacing = 1.15
        cap.runs[0].italic = True


def ref(doc: Document, text: str):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.line_spacing = 1.0
    para.add_run(text)


def get_metric(model: str, split: str, col: str = "auroc_mean") -> float:
    row = bench[(bench["model"] == model) & (bench["split"] == split)]
    return float(row.iloc[0][col])


def build_markdown() -> str:
    return f"""# Completed pilot manuscript v1

This draft includes a completed Phase 0 metadata audit and a blood T cell embedding benchmark.

Key pilot result: dataset-holdout performance was near chance for scVI ({get_metric('scVI_embedding_logistic','dataset_holdout'):.3f}) and TranscriptFormer ({get_metric('TranscriptFormer_embedding_logistic','dataset_holdout'):.3f}), while donor-holdout retained moderate signal.
"""


def build_docx() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Donor-, disease-, and study-aware benchmarking of single-cell embeddings for human cellular aging")
    r.bold = True
    r.font.size = Pt(16)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("First complete research manuscript draft with metadata audit and blood T cell pilot benchmark").italic = True
    p(doc, "Author information: to be completed before external submission.")
    p(doc, "Correspondence: to be completed before external submission.")

    doc.add_heading("Abstract", level=1)
    p(doc, "Background: Single-cell foundation models and integrated embeddings are increasingly used to represent cellular states, but their apparent performance in aging prediction may be inflated by donor overlap, disease imbalance, tissue composition, study origin, and assay effects. A leakage-aware benchmark is needed before such models can be interpreted as learning transferable cellular aging biology.")
    p(doc, f"Methods: We first performed a CELLxGENE Census metadata feasibility audit across five target cell systems using release {meta_summary['census_version']}. We then ran a minimum complete benchmark in blood T cells, the strongest eligible stratum with numeric age information. Young and old donors were defined using numeric development_stage-derived ages of <=40 and >=60 years. The pilot used {int(bench_audit['sample_all_donors'])} donors ({int(bench_audit['sample_all_young_donors'])} young and {int(bench_audit['sample_all_old_donors'])} old) and {int(bench_audit['sample_all_cells'])} sampled cells. Models included a metadata-confounder baseline, CELLxGENE-hosted scVI embeddings, CELLxGENE-hosted TranscriptFormer embeddings, and donor-level mean-embedding pseudobulk models. We evaluated random cell split, donor-holdout, dataset-holdout, disease-free donor-holdout, and label-shuffling controls.")
    p(doc, f"Results: The metadata audit covered {int(meta_summary['target_primary_cell_rows']):,} primary-cell metadata rows from {int(meta_summary['unique_donors']):,} donors and identified 18 eligible tissue-cell strata under the primary threshold. In the blood T cell benchmark, random-split AUROC was high for the metadata baseline ({get_metric('metadata_confounder_logistic','random_cell'):.3f}), TranscriptFormer embeddings ({get_metric('TranscriptFormer_embedding_logistic','random_cell'):.3f}), and scVI embeddings ({get_metric('scVI_embedding_logistic','random_cell'):.3f}). Dataset-holdout performance dropped to near-chance or below-chance levels for TranscriptFormer ({get_metric('TranscriptFormer_embedding_logistic','dataset_holdout'):.3f}) and scVI ({get_metric('scVI_embedding_logistic','dataset_holdout'):.3f}), whereas donor-holdout retained moderate signal. Label shuffling reduced random-split AUROC toward chance, producing permutation gaps of 0.263-0.288 across the three cell-level models.")
    p(doc, "Conclusions: Public single-cell aging prediction is feasible in selected strata, but apparent performance is strongly shaped by metadata and study structure. The completed pilot supports a benchmark strategy that treats random-split accuracy as an insufficient endpoint and prioritizes donor-holdout, dataset-holdout, disease-free, and permutation controls.")

    doc.add_heading("Keywords", level=1)
    p(doc, "single-cell RNA sequencing; cellular aging; foundation models; scVI; TranscriptFormer; donor holdout; dataset holdout; disease confounding; CELLxGENE Census")

    doc.add_heading("Background", level=1)
    p(doc, "Molecular aging predictors have shown that chronological and biological age can be estimated from molecular profiles. DNA methylation clocks established this principle across tissues and cell types, and transcriptomic resources now make cell-type-specific aging analyses possible. Single-cell transcriptomics adds a distinct opportunity: the same tissue can be decomposed into cell types and cellular states rather than analyzed as a bulk mixture.")
    p(doc, "The same resolution also creates evaluation risks. In public single-cell atlases, older and younger donors may come from different studies, disease cohorts, tissues, assays, or recruitment settings. A model can therefore appear to predict age while learning donor identity, disease status, study origin, or cell-composition structure. This is particularly relevant for foundation-model embeddings, whose representations are flexible enough to capture both biology and confounding structure.")
    p(doc, "We designed this study to separate metadata feasibility, apparent prediction, and leakage-resistant generalization. The first stage audits whether public data contain enough donor, age, disease, and dataset structure for a fair benchmark. The second stage tests a minimum benchmark in a stratum that passes this audit. The primary endpoint is not the highest random-split accuracy, but the extent to which performance persists beyond donor, disease, and study structure.")

    doc.add_heading("Methods", level=1)
    doc.add_heading("Metadata feasibility audit", level=2)
    p(doc, f"The Phase 0 audit used CELLxGENE Census release {meta_summary['census_version']}. We queried primary human-cell metadata for T cells, monocytes, macrophages, endothelial cells, and fibroblasts. A tissue-cell stratum was eligible when it contained at least 30 donors, at least 10 young/adult donors, at least 10 old/elderly donors, at least two datasets, and disease-free donor retention of at least 0.50. Threshold sensitivity was evaluated by varying donor, age-proxy, and disease-free requirements.")
    p(doc, "Development_stage was used for feasibility screening. For the completed pilot benchmark, only blood T cells with numeric year-old or decade-derived ages were used, and young-old classification was restricted to donors <=40 or >=60 years. Middle-aged donors were excluded from the classification task.")

    doc.add_heading("Blood T cell pilot benchmark", level=2)
    p(doc, f"The blood T cell stratum contained {int(bench_audit['full_blood_t_cell_metadata_rows']):,} metadata rows, of which {int(bench_audit['numeric_young_old_metadata_rows']):,} belonged to numeric young or old age groups. We sampled {int(bench_audit['sample_all_donors'])} donors balanced by age group and up to eight cells per donor, yielding {int(bench_audit['sample_all_cells'])} cells. A disease-free subset was sampled separately with {int(bench_audit['sample_disease_free_donors'])} donors and {int(bench_audit['sample_disease_free_cells'])} cells.")
    p(doc, "The model set included three cell-level models and two donor-level models. The metadata-confounder baseline used dataset_id, assay, disease, and sex. The scVI model used CELLxGENE-hosted 50-dimensional scVI embeddings. The TranscriptFormer model used CELLxGENE-hosted 2048-dimensional TF-Sapiens embeddings. Donor-level pseudobulk analogues were created by averaging embeddings within each donor and fitting the same logistic model.")
    p(doc, "All predictive models used logistic regression with class balancing. Evaluations were repeated five times for random cell split, donor-holdout split, and dataset-holdout split. The disease-free subset was evaluated under the same split families. Label-shuffling controls permuted training labels while preserving the test set and data structure.")

    doc.add_heading("Statistical analysis", level=2)
    p(doc, "The main classification metrics were AUROC, AUPRC, and balanced accuracy. Generalization drop was defined as random-split AUROC minus donor-holdout AUROC. Permutation gap was defined as true-label random-split AUROC minus shuffled-label random-split AUROC. Results are reported as means across five repeated splits with standard deviations in the supplementary tables.")

    doc.add_heading("Results", level=1)
    doc.add_heading("Metadata audit supported a restricted benchmark rather than pan-tissue assumptions", level=2)
    p(doc, f"The Phase 0 audit retrieved {int(meta_summary['target_primary_cell_rows']):,} primary-cell metadata rows across five target cell systems, representing {int(meta_summary['unique_donors']):,} donors, {int(meta_summary['unique_datasets']):,} datasets, {int(meta_summary['unique_collections']):,} collections, and {int(meta_summary['unique_tissues']):,} tissues. It assessed {len(tissue)} tissue-cell strata. Eighteen strata met the primary eligibility threshold, and 98 met pilot-candidate criteria.")
    pic(doc, "outputs/figures_v2/figure1_metadata_screening_flow.png", "Figure 1. Metadata screening flow for the Phase 0 audit.")
    p(doc, "T cells, macrophages, and fibroblasts met system-level eligibility criteria. Endothelial cells and monocytes had broad donor and dataset coverage but disease-free retention below the primary threshold, so they were retained as pilot candidates.")
    pic(doc, "outputs/figures_v2/figure4_disease_free_ci.png", "Figure 2. Disease-free donor retention by target cell system with 95% Wilson confidence intervals.")
    base = sens[(sens["donor_min"] == 30) & (sens["age_group_min"] == 10) & (sens["disease_free_min"] == 0.5)].iloc[0]
    relaxed = sens[(sens["donor_min"] == 30) & (sens["age_group_min"] == 10) & (sens["disease_free_min"] == 0.3)].iloc[0]
    strict = sens[(sens["donor_min"] == 40) & (sens["age_group_min"] == 10) & (sens["disease_free_min"] == 0.6)].iloc[0]
    p(doc, f"Threshold sensitivity supported the same qualitative conclusion. Under the primary threshold, {int(base['eligible_strata_n'])} tissue-cell strata were eligible. Relaxing the disease-free retention threshold to 0.30 yielded {int(relaxed['eligible_strata_n'])} eligible strata, whereas a stricter rule requiring at least 40 donors and disease-free retention of 0.60 yielded {int(strict['eligible_strata_n'])} eligible strata.")
    pic(doc, "outputs/figures_v2/figure3_threshold_sensitivity.png", "Figure 3. Threshold sensitivity analysis across donor and disease-free retention thresholds.")

    doc.add_heading("Random split performance was high but metadata alone was predictive", level=2)
    p(doc, f"In blood T cells, random cell split AUROC was {get_metric('metadata_confounder_logistic','random_cell'):.3f} for the metadata-confounder baseline, {get_metric('TranscriptFormer_embedding_logistic','random_cell'):.3f} for TranscriptFormer embeddings, and {get_metric('scVI_embedding_logistic','random_cell'):.3f} for scVI embeddings. The high metadata-baseline performance indicates that dataset, assay, disease, and sex carried age-predictive structure in the sampled public data.")

    doc.add_heading("Dataset holdout revealed study-level fragility", level=2)
    p(doc, f"Dataset-holdout evaluation caused the strongest performance loss. AUROC dropped to {get_metric('TranscriptFormer_embedding_logistic','dataset_holdout'):.3f} for TranscriptFormer embeddings and {get_metric('scVI_embedding_logistic','dataset_holdout'):.3f} for scVI embeddings. The metadata baseline dropped to {get_metric('metadata_confounder_logistic','dataset_holdout'):.3f}, with high split-to-split variability. These results suggest that much of the apparent age-prediction signal is aligned with study structure rather than being uniformly transferable across datasets.")
    p(doc, f"Donor-holdout performance was more stable than dataset-holdout performance. TranscriptFormer and scVI achieved donor-holdout AUROCs of {get_metric('TranscriptFormer_embedding_logistic','donor_holdout'):.3f} and {get_metric('scVI_embedding_logistic','donor_holdout'):.3f}, respectively. Donor-level mean-embedding models performed better under donor-holdout, with AUROC {get_metric('TranscriptFormer_mean_embedding_pseudobulk','donor_holdout'):.3f} for TranscriptFormer and {get_metric('scVI_mean_embedding_pseudobulk','donor_holdout'):.3f} for scVI.")
    pic(doc, "outputs/benchmark_figures/figure_blood_tcell_embedding_performance.png", "Figure 4. Blood T cell young-old prediction under random, donor-holdout, dataset-holdout, and disease-free donor-holdout splits.")

    doc.add_heading("Disease-free evaluation retained embedding signal while reducing metadata confounding", level=2)
    p(doc, f"In the disease-free donor-holdout subset, TranscriptFormer and scVI cell-level embeddings achieved AUROCs of {get_metric('TranscriptFormer_embedding_logistic','disease_free_donor_holdout'):.3f} and {get_metric('scVI_embedding_logistic','disease_free_donor_holdout'):.3f}, respectively. The metadata baseline was lower at {get_metric('metadata_confounder_logistic','disease_free_donor_holdout'):.3f}. Donor-level mean embeddings showed stronger disease-free donor-holdout performance, with AUROCs of {get_metric('TranscriptFormer_mean_embedding_pseudobulk','disease_free_donor_holdout'):.3f} and {get_metric('scVI_mean_embedding_pseudobulk','disease_free_donor_holdout'):.3f}.")

    doc.add_heading("Permutation controls supported signal above shuffled-label baselines", level=2)
    p(doc, "Label-shuffling controls reduced random-split performance toward chance for all three cell-level models. The permutation gap was 0.288 for TranscriptFormer embeddings, 0.286 for the metadata baseline, and 0.263 for scVI embeddings. The donor-holdout generalization drop from random split was smaller, ranging from 0.011 for scVI to 0.033 for TranscriptFormer, indicating that donor overlap was not the only source of apparent performance.")
    pic(doc, "outputs/benchmark_figures/figure_blood_tcell_embedding_diagnostics.png", "Figure 5. Generalization drop and permutation gap for blood T cell cell-level models.")

    doc.add_paragraph("Table 1. Summary of selected benchmark AUROC results.")
    selected = bench[
        bench["model"].isin([
            "metadata_confounder_logistic",
            "scVI_embedding_logistic",
            "TranscriptFormer_embedding_logistic",
            "scVI_mean_embedding_pseudobulk",
            "TranscriptFormer_mean_embedding_pseudobulk",
        ])
        & bench["split"].isin(["random_cell", "donor_holdout", "dataset_holdout", "disease_free_donor_holdout"])
    ].copy()
    model_labels = {
        "metadata_confounder_logistic": "Metadata",
        "scVI_embedding_logistic": "scVI cell",
        "TranscriptFormer_embedding_logistic": "TF cell",
        "scVI_mean_embedding_pseudobulk": "scVI donor mean",
        "TranscriptFormer_mean_embedding_pseudobulk": "TF donor mean",
    }
    split_labels = {
        "random_cell": "Random",
        "donor_holdout": "Donor holdout",
        "dataset_holdout": "Dataset holdout",
        "disease_free_donor_holdout": "Disease-free donor",
    }
    selected["Model"] = selected["model"].map(model_labels).fillna(selected["model"])
    selected["Split"] = selected["split"].map(split_labels).fillna(selected["split"])
    selected["AUROC"] = selected["auroc_mean"].round(3)
    selected["AUPRC"] = selected["auprc_mean"].round(3)
    selected["Balanced acc."] = selected["balanced_accuracy_mean"].round(3)
    table(doc, selected[["Model", "Split", "AUROC", "AUPRC", "Balanced acc."]], ["Model", "Split", "AUROC", "AUPRC", "Balanced acc."])

    doc.add_heading("Discussion", level=1)
    p(doc, "This first complete pilot benchmark shows why single-cell aging prediction needs leakage-aware evaluation. Random-split performance was not sufficient evidence of transferable aging biology, because metadata alone predicted young-old status with AUROC above 0.77. This result is a warning: public single-cell aging benchmarks can be driven by study, assay, disease, or sex structure unless these variables are explicitly audited.")
    p(doc, "The strongest stress test was dataset holdout. Both scVI and TranscriptFormer embeddings performed near chance or below chance when entire datasets were held out, despite retaining moderate donor-holdout performance. This pattern suggests that donor-level separation alone is necessary but not sufficient. Study-level differences remain a major barrier to generalization.")
    p(doc, "The disease-free donor-holdout results were more encouraging. Embedding-based models retained signal after restricting to disease-free donors, while the metadata baseline weakened. Donor-level mean embeddings performed particularly well in this pilot, suggesting that averaging across cells may reduce cell-level noise and emphasize donor-level age-associated programs. This finding should be treated as hypothesis-generating until replicated across additional eligible strata.")
    p(doc, "The study has several limitations. First, the complete benchmark currently covers only blood T cells. Second, numeric age was recovered from development_stage labels rather than a harmonized donor-age field. Third, TranscriptFormer is represented by hosted embeddings from a current Census release, while Geneformer and scGPT were not run directly in this pilot. Fourth, dataset-holdout estimates had high split-to-split variability because some held-out studies carried strong age and disease structure. These limitations motivate expansion to fibroblast and macrophage strata and a more explicit source-metadata recovery step.")
    p(doc, "Despite these limitations, the pilot changes the manuscript from a feasibility report into an empirical benchmark. The key result is not that one embedding is universally superior, but that apparent cellular aging prediction is strongly dependent on how the split is constructed. A credible benchmark should therefore report random split, donor-holdout, dataset-holdout, disease-free, and label-shuffling results together.")

    doc.add_heading("Conclusions", level=1)
    p(doc, "Public single-cell aging prediction is feasible in selected human cell strata, but model performance is strongly shaped by metadata and study structure. In blood T cells, hosted scVI and TranscriptFormer embeddings retained moderate donor-holdout and disease-free donor-holdout signal, whereas dataset-holdout performance was near chance. These results support a donor-, disease-, and study-aware benchmark framework and argue against interpreting random-split accuracy as evidence of transferable cellular aging biology.")

    doc.add_heading("Supplementary Information", level=1)
    p(doc, "Supplementary tables include the metadata feasibility map, threshold sensitivity analysis, blood T cell benchmark long-format results, summary metrics, leakage diagnostics, and benchmark audit JSON. These files are included in the local project directory and should be deposited in a public repository before external submission.")
    doc.add_heading("Acknowledgements", level=1)
    p(doc, "To be completed before submission.")
    doc.add_heading("Authors' contributions", level=1)
    p(doc, "To be completed before submission.")
    doc.add_heading("Funding", level=1)
    p(doc, "To be completed before submission.")
    doc.add_heading("Availability of data and materials", level=1)
    p(doc, "The analyses used CELLxGENE Census release 2025-11-08 and hosted embeddings from the same release. Current scripts and generated summary tables are stored in this project directory. Before submission, the code and non-sensitive summary tables should be archived through GitHub and Zenodo or an equivalent repository.")
    doc.add_heading("Ethics approval and consent to participate", level=1)
    p(doc, "This study used public de-identified single-cell metadata and hosted embeddings. No new human participants or animal experiments were included.")
    doc.add_heading("Consent for publication", level=1)
    p(doc, "Not applicable.")
    doc.add_heading("Competing interests", level=1)
    p(doc, "To be completed before submission.")

    doc.add_heading("References", level=1)
    refs = [
        "Horvath S. DNA methylation age of human tissues and cell types. Genome Biology. 2013;14:R115. doi:10.1186/gb-2013-14-10-r115.",
        "Hannum G, Guinney J, Zhao L, Zhang L, Hughes G, Sadda S, et al. Genome-wide methylation profiles reveal quantitative views of human aging rates. Molecular Cell. 2013;49:359-367. doi:10.1016/j.molcel.2012.10.016.",
        "Lopez R, Regier J, Cole MB, Jordan MI, Yosef N. Deep generative modeling for single-cell transcriptomics. Nature Methods. 2018;15:1053-1058. doi:10.1038/s41592-018-0229-2.",
        "Theodoris CV, Xiao L, Chopra A, Chaffin MD, Al Sayed ZR, Hill MC, et al. Transfer learning enables predictions in network biology. Nature. 2023;618:616-624. doi:10.1038/s41586-023-06139-9.",
        "Cui H, Wang C, Maan H, Pang K, Luo F, Duan N, et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods. 2024. doi:10.1038/s41592-024-02201-0.",
        "Bartz J, Ma X, Zhang L, Dong X. Human Cell Aging Transcriptome Atlas (HCATA): a single-cell atlas of age-associated transcriptomic alterations across human tissues. Communications Biology. 2025. doi:10.1038/s42003-025-08845-8.",
        "Aging Atlas: a multi-omics database for aging biology. Nucleic Acids Research. 2021. doi:10.1093/nar/gkaa894.",
        "The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science. 2020;369:1318-1330. doi:10.1126/science.aaz1776.",
        "Zou Z, Long X, Zhao Q, Zheng Y, Song M, Ma S, et al. A single-cell transcriptomic atlas of human skin aging. Developmental Cell. 2021;56:383-397.e8. doi:10.1016/j.devcel.2020.11.002.",
        "Han X, Zhou Z, Fei L, Sun H, Wang R, Chen Y, et al. Construction of a human cell landscape at single-cell level. Nature. 2020;581:303-309. doi:10.1038/s41586-020-2157-4.",
        "CELLxGENE Census documentation. CZ CELLxGENE Discover Census and hosted embeddings documentation. Available from: https://chanzuckerberg.github.io/cellxgene-census/.",
        "TranscriptFormer TF-Sapiens cell embeddings. CELLxGENE Census hosted embedding metadata, release 2025-11-08. Associated preprint doi:10.1101/2025.04.25.650731.",
    ]
    for item in refs:
        ref(doc, item)

    for section in doc.sections:
        add_line_numbers(section)
    doc.save(DOCX)
    MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
