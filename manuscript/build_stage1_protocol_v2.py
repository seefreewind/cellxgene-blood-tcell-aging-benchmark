from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(".")
OUT = ROOT / "outputs"
DOCX = OUT / "donor_disease_aware_sc_aging_stage1_protocol_v2.docx"
MD = ROOT / "manuscript" / "donor_disease_aware_sc_aging_stage1_protocol_v2.md"

summary = pd.read_json(ROOT / "analysis/outputs/metadata_audit_summary.json", typ="series")
system = pd.read_csv(ROOT / "analysis/outputs/cell_system_feasibility_with_ci.tsv", sep="\t")
tissue = pd.read_csv(ROOT / "analysis/outputs/metadata_feasibility_map.tsv", sep="\t")
sensitivity = pd.read_csv(ROOT / "analysis/outputs/threshold_sensitivity.tsv", sep="\t")
age_proxy = pd.read_csv(ROOT / "analysis/outputs/age_proxy_distribution_by_cell_system.tsv", sep="\t")
sex_dist = pd.read_csv(ROOT / "analysis/outputs/sex_distribution_by_cell_system.tsv", sep="\t")
flow = pd.read_csv(ROOT / "analysis/outputs/metadata_screening_flow.tsv", sep="\t")


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "newPage")
    sect_pr.append(ln)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    for i, col in enumerate(columns):
        set_cell(table.rows[0].cells[i], col, True)
        shade(table.rows[0].cells[i], "E8EEF5")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = f"{value:.3f}" if abs(value) < 10 else f"{value:,.0f}"
            elif isinstance(value, int):
                value = f"{value:,}"
            set_cell(cells[i], str(value))
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: str, caption: str, width: float = 6.2) -> None:
    p = ROOT / path
    if p.exists():
        doc.add_picture(str(p), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True


def add_ref(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.0
    p.add_run(text)


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_line_numbers(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = False
        style.paragraph_format.line_spacing = 1.15


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 2.0


def build_markdown() -> str:
    base = sensitivity[
        (sensitivity["donor_min"] == 30)
        & (sensitivity["age_group_min"] == 10)
        & (sensitivity["disease_free_min"] == 0.5)
    ].iloc[0]
    relaxed = sensitivity[
        (sensitivity["donor_min"] == 30)
        & (sensitivity["age_group_min"] == 10)
        & (sensitivity["disease_free_min"] == 0.3)
    ].iloc[0]
    return f"""# Stage 1 protocol v2 summary

Primary repositioning: the document is now framed as a registered-report-style Stage 1 benchmark protocol with completed Phase 0 metadata feasibility results.

Key completed results:

- CELLxGENE Census release: {summary['census_version']}
- Target primary-cell metadata rows: {int(summary['target_primary_cell_rows']):,}
- Donors: {int(summary['unique_donors']):,}
- Datasets: {int(summary['unique_datasets']):,}
- Tissue-cell strata assessed: {len(tissue):,}
- Eligible tissue-cell strata under the primary threshold: {int(base['eligible_strata_n'])}
- Eligible tissue-cell strata under relaxed disease-free threshold 0.30: {int(relaxed['eligible_strata_n'])}

Main limitation retained explicitly: numeric age and model-performance analyses remain required before converting this document into a conventional research article.
"""


def build_docx() -> None:
    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "Donor- and disease-aware benchmarking of single-cell foundation models for human cellular aging"
    )
    r.bold = True
    r.font.size = Pt(16)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Registered-report-style Stage 1 protocol with completed metadata feasibility audit").italic = True
    add_paragraph(doc, "Manuscript type: Stage 1 benchmark protocol / methods protocol.")
    add_paragraph(doc, "Author information: to be completed before external submission.")
    add_paragraph(doc, "Correspondence: to be completed before external submission.")

    doc.add_heading("Abstract", level=1)
    add_paragraph(doc, "Background: Single-cell foundation models are increasingly used to represent cellular states, but their apparent performance in aging prediction can be inflated by donor overlap, disease imbalance, study origin, tissue composition, and assay effects. Existing molecular aging work, including DNA methylation clocks, shows that biological age can be predicted from molecular profiles, but single-cell transcriptomic benchmarks face additional compositional and metadata-leakage risks.")
    add_paragraph(doc, f"Methods: We designed a registered-report-style benchmark protocol in which metadata feasibility is the first decision gate. Using CELLxGENE Census release {summary['census_version']}, we audited primary human-cell metadata for T cells, monocytes, macrophages, endothelial cells, and fibroblasts. Eligibility was evaluated at cell-system and tissue-cell levels using prespecified donor, age-proxy, dataset, and disease-free retention thresholds. We also quantified threshold sensitivity, Wilson confidence intervals for disease-free donor retention, and age-proxy and sex distributions.")
    add_paragraph(doc, f"Results: The Phase 0 audit included {int(summary['target_primary_cell_rows']):,} primary-cell metadata rows from {int(summary['unique_donors']):,} donors, {int(summary['unique_datasets']):,} datasets, {int(summary['unique_collections']):,} collections, and {int(summary['unique_tissues']):,} tissues. T cells, macrophages, and fibroblasts met the system-level eligibility criteria, whereas endothelial cells and monocytes were retained as pilot candidates because disease-free donor retention was below 0.50. At the tissue-cell level, 18 of {len(tissue)} assessed strata met the primary eligibility threshold. Threshold sensitivity showed that relaxing the disease-free threshold from 0.50 to 0.30 increased eligible strata from 18 to 21 under the same donor and age-proxy requirements.")
    add_paragraph(doc, "Conclusions: Public single-cell aging benchmarks should not begin from an assumption of pan-tissue comparability. The completed metadata audit supports an initial leakage-resistant benchmark in selected immune and stromal strata, with numeric age recovery, donor-holdout, study-holdout, disease-free, and label-shuffling analyses required before claims about transferable cellular aging biology can be made.")

    doc.add_heading("Keywords", level=1)
    add_paragraph(doc, "single-cell RNA sequencing; cellular aging; foundation models; registered report; donor holdout; disease confounding; CELLxGENE Census; benchmark")

    doc.add_heading("Background", level=1)
    add_paragraph(doc, "Molecular aging prediction has a long precedent in bulk molecular data. DNA methylation clocks demonstrated that chronological and biological age can be estimated across tissues and cell types, and later methylation-based predictors provided quantitative views of human aging rates. These approaches established the conceptual value of molecular age prediction, but they do not remove the need for leakage-aware evaluation in single-cell transcriptomics, where cell composition, donor identity, disease status, and study design can be entangled with age.")
    add_paragraph(doc, "Public single-cell atlases now contain enough human-cell data to make cellular aging benchmarks technically plausible. CELLxGENE Discover and CELLxGENE Census provide standardized programmatic access to large-scale single-cell data, while HCATA and Aging Atlas provide aging-oriented reference resources for interpretation. These resources make cross-study analyses possible, but they also aggregate datasets that differ in donor recruitment, disease composition, tissue sampling, assay, and annotation depth.")
    add_paragraph(doc, "Foundation models such as Geneformer and scGPT offer reusable transcriptomic representations and may encode cellular programs relevant to aging. However, a representation that performs well under random cell splits may not generalize when donors, studies, or disease states are held out. For aging, this distinction is central: a model may appear to predict age because older cells come from particular donors, disease cohorts, tissues, or datasets.")
    add_paragraph(doc, "We therefore frame this work as a staged, leakage-aware benchmark protocol. The first stage asks whether public metadata support strict evaluation at all. The second stage, to be run only in eligible strata, will test whether model performance persists under donor-holdout, study-holdout, disease-free, and label-shuffling controls. The primary endpoint is not maximal random-split accuracy, but leakage-resistant generalization beyond donor, disease, and study structure.")

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design and decision gates", level=2)
    add_paragraph(doc, "This manuscript is structured as a Stage 1 benchmark protocol with completed metadata feasibility results. The benchmark follows three gates. Gate 1 is metadata feasibility: a tissue-cell stratum must have adequate donor, age-proxy, dataset, and disease-free coverage. Gate 2 is numeric-age recovery: strata that pass Gate 1 must be linked back to source metadata to recover numeric donor age or age ranges. Gate 3 is leakage-resistant performance: models must be evaluated under donor-holdout, study-holdout, disease-free, and label-shuffling controls.")
    add_paragraph(doc, "A conventional research-article version of this manuscript will require Gate 2 and Gate 3 results. Until those analyses are complete, conclusions are restricted to feasibility and protocol-level claims.")

    doc.add_heading("Metadata audit", level=2)
    add_paragraph(doc, f"The Phase 0 audit used CELLxGENE Census release {summary['census_version']}. We queried primary human-cell metadata for five target cell systems: T cells, monocytes, macrophages, endothelial cells, and fibroblasts. Retrieved fields included dataset identifier, collection identifier, assay, tissue, tissue_general, cell_type, disease, sex, development_stage, donor_id, suspension_type, and is_primary_data. Dataset-level collection metadata were joined from the Census dataset table.")
    add_paragraph(doc, "Because numeric donor age was not available in the queried Census obs schema, development_stage was used only as an age-proxy field. Labels were classified into young_or_adult, old_or_elderly, annotated_other, or unknown. This proxy is acceptable for feasibility screening, but it is not sufficient for the final age-prediction benchmark. In the next stage, exact donor age or age-range midpoints must be recovered from source datasets wherever permitted.")
    add_paragraph(doc, "A tissue-cell stratum was eligible when it contained at least 30 donors, at least 10 young/adult donors, at least 10 old/elderly donors, at least two datasets, and disease-free donor retention of at least 0.50. Strata were pilot candidates when they contained at least 15 donors, at least two datasets, and disease-free donor retention of at least 0.30. These thresholds were selected to avoid single-study, single-age-group, or disease-dominated benchmarks and were evaluated in a sensitivity analysis.")

    doc.add_heading("Planned benchmark after numeric-age recovery", level=2)
    add_paragraph(doc, "The minimum publishable benchmark will include blood T cells and at least one stromal stratum, with macrophages or monocytes and endothelial cells included as pilot or sensitivity strata where metadata permit. Primary models will include pseudobulk ridge or elastic net, PCA plus linear or logistic regression, scVI embeddings, and hosted or released foundation-model embeddings. De novo large-scale foundation-model inference will not be the default unless hosted embeddings are insufficient.")
    add_paragraph(doc, "The primary evaluation splits will be random cell split, donor-holdout, dataset or study holdout, disease-free only evaluation, disease-stratified evaluation where feasible, and label-shuffling or permutation baselines that preserve donor, tissue, and dataset structure. A disease-prediction negative control will be used to test whether representations encode disease status more strongly than leakage-resistant age signal.")
    add_paragraph(doc, "For continuous age prediction, planned metrics are mean absolute error, Spearman correlation, and R-squared. For young-old classification, planned metrics are AUROC, AUPRC, and balanced accuracy. Generalization drop will be defined as random-split performance minus holdout performance. Permutation gap will be defined as true-label performance minus shuffled-label performance.")

    doc.add_heading("Results", level=1)
    doc.add_heading("Metadata screening identified a restricted set of benchmark-ready strata", level=2)
    add_paragraph(doc, f"The Phase 0 audit retrieved {int(summary['target_primary_cell_rows']):,} primary-cell metadata rows, representing {int(summary['unique_donors']):,} donors, {int(summary['unique_datasets']):,} datasets, {int(summary['unique_collections']):,} collections, and {int(summary['unique_tissues']):,} tissues. The screening flow reduced these resources to {len(tissue)} assessed tissue-cell strata, of which 18 met the primary eligibility threshold and 98 met pilot-candidate criteria.")
    add_picture_if_exists(doc, "outputs/figures_v2/figure1_metadata_screening_flow.png", "Figure 1. Metadata screening flow for the Phase 0 audit. The figure shows the transition from target CELLxGENE metadata rows to assessed tissue-cell strata and eligible or pilot-candidate strata.")

    doc.add_heading("System-level eligibility favored immune and stromal systems", level=2)
    add_paragraph(doc, "T cells, macrophages, and fibroblasts met the system-level eligibility criteria. T cells included 3,791 donors and 1,267,019 cells, with disease-free donor retention of 0.548 (95% Wilson CI, 0.532-0.564). Macrophages included 2,984 donors and 1,736,309 cells, with disease-free retention of 0.504 (95% CI, 0.486-0.522). Fibroblasts included 2,214 donors and 2,663,513 cells, with disease-free retention of 0.586 (95% CI, 0.566-0.607). Endothelial cells and monocytes had broad donor and dataset coverage but disease-free retention below the primary 0.50 threshold.")
    add_picture_if_exists(doc, "outputs/figures_v2/figure4_disease_free_ci.png", "Figure 2. Disease-free donor retention by cell system with 95% Wilson confidence intervals. The dashed line marks the primary 0.50 disease-free retention threshold.")

    table_df = system.copy()
    table_df["Disease-free"] = table_df["disease_free_donor_fraction"].round(3)
    table_df["95% CI"] = table_df.apply(lambda r: f"{r['disease_free_ci_low']:.3f}-{r['disease_free_ci_high']:.3f}", axis=1)
    table_df["Status"] = table_df["feasibility_status"].map({"eligible_stratum": "eligible", "pilot_candidate": "pilot"}).fillna(table_df["feasibility_status"])
    table_df = table_df.rename(columns={"cell_type": "Cell system", "donor_n": "Donors", "cell_n": "Cells", "tissue_n": "Tissues", "dataset_n": "Datasets"})
    table_df["Cell system"] = table_df["Cell system"].replace({"endothelial cell": "endothelial"})
    doc.add_page_break()
    add_df_table(doc, table_df[["Cell system", "Donors", "Cells", "Tissues", "Datasets", "Disease-free", "95% CI", "Status"]], ["Cell system", "Donors", "Cells", "Tissues", "Datasets", "Disease-free", "95% CI", "Status"])

    doc.add_heading("Threshold sensitivity showed that eligibility was not driven by a single arbitrary cutoff", level=2)
    base = sensitivity[(sensitivity["donor_min"] == 30) & (sensitivity["age_group_min"] == 10) & (sensitivity["disease_free_min"] == 0.5)].iloc[0]
    relaxed = sensitivity[(sensitivity["donor_min"] == 30) & (sensitivity["age_group_min"] == 10) & (sensitivity["disease_free_min"] == 0.3)].iloc[0]
    strict = sensitivity[(sensitivity["donor_min"] == 40) & (sensitivity["age_group_min"] == 10) & (sensitivity["disease_free_min"] == 0.6)].iloc[0]
    add_paragraph(doc, f"Under the primary threshold, {int(base['eligible_strata_n'])} tissue-cell strata were eligible. Relaxing the disease-free retention threshold to 0.30 while keeping the same donor and age-proxy requirements yielded {int(relaxed['eligible_strata_n'])} eligible strata. A stricter rule requiring at least 40 donors and disease-free retention of 0.60 yielded {int(strict['eligible_strata_n'])} eligible strata. Thus, the main conclusion, that eligibility is concentrated in selected immune and stromal strata rather than pan-tissue coverage, was stable across reasonable threshold changes.")
    add_picture_if_exists(doc, "outputs/figures_v2/figure3_threshold_sensitivity.png", "Figure 3. Threshold sensitivity analysis. Eligible tissue-cell strata are shown across donor and disease-free retention thresholds while holding the minimum age-proxy group size at 10 donors.")

    doc.add_heading("Age-proxy and sex summaries define the next metadata recovery task", level=2)
    add_paragraph(doc, "Age-proxy summaries showed that all five target systems contained old_or_elderly donors, but young_or_adult coverage was more limited in endothelial cells and monocytes. Sex metadata were available in the donor-level summaries and should be used for stratified reporting in the final benchmark. These summaries identify the next bottleneck: exact numeric donor age must be recovered from source datasets before age regression can be treated as the primary endpoint.")
    add_picture_if_exists(doc, "outputs/figures_v2/figure5_age_proxy_distribution.png", "Figure 4. Age-proxy composition by target cell system. Development_stage-derived categories are used only for feasibility screening and will not replace numeric age in the final benchmark.")

    age_summary = age_proxy.pivot_table(index="cell_type", columns="age_group_proxy", values="donor_n", fill_value=0).reset_index()
    age_summary = age_summary.rename(columns={"cell_type": "Cell system", "young_or_adult": "Young/adult", "old_or_elderly": "Old/elderly", "annotated_other": "Other"})
    for col in ["Young/adult", "Old/elderly", "Other", "unknown"]:
        if col not in age_summary.columns:
            age_summary[col] = 0
    add_df_table(doc, age_summary[["Cell system", "Young/adult", "Old/elderly", "Other"]], ["Cell system", "Young/adult", "Old/elderly", "Other"])

    doc.add_heading("Discussion", level=1)
    add_paragraph(doc, "This revised draft changes the manuscript's center of gravity. The current evidence does not yet support claims about foundation-model performance in aging prediction, but it does support a rigorous Stage 1 benchmark protocol and a completed metadata feasibility audit. This distinction improves the manuscript because it treats missing model results as a design boundary rather than a weakness hidden inside the Results section.")
    add_paragraph(doc, "The audit indicates that a pan-tissue benchmark should not be assumed. Eligibility was strongest for T cells, macrophages, and fibroblasts, while endothelial cells and monocytes required pilot handling because disease-free retention was lower. This is not merely a data-volume issue. Disease-free retention, age-proxy balance, and dataset diversity determine whether a benchmark can distinguish aging signal from disease, donor, and study structure.")
    add_paragraph(doc, "The next version should move from Stage 1 protocol to conventional benchmark article only after numeric donor age and expression or embedding data have been analyzed. The decisive figure will be a performance-collapse plot comparing random, donor-holdout, study-holdout, disease-free, and shuffled-label evaluations. If random-split performance is high but holdout performance collapses, the paper can make a strong leakage-audit claim. If some models retain performance under strict controls, the paper can then interpret robust aging programs.")
    add_paragraph(doc, "The main limitation is that development_stage is only an age proxy. It is useful for screening, but it cannot support final claims about chronological age prediction. A second limitation is that disease-free status is inferred from normal or healthy labels, which may not capture donor-level clinical detail. These limitations define the required next analyses rather than invalidating the metadata audit.")

    doc.add_heading("Conclusions", level=1)
    add_paragraph(doc, "The completed metadata audit supports a restricted, leakage-aware benchmark of human cellular aging in selected immune and stromal strata. The manuscript should remain a Stage 1 protocol until numeric age recovery and model-performance analyses are complete. The planned benchmark should prioritize donor-holdout, study-holdout, disease-free, disease-negative-control, and label-shuffling analyses over random-split accuracy.")

    doc.add_heading("Supplementary Information", level=1)
    add_paragraph(doc, "Supplementary Table 1: cell-system feasibility with Wilson confidence intervals. Supplementary Table 2: tissue-cell feasibility map. Supplementary Table 3: threshold sensitivity analysis. Supplementary Table 4: age-proxy and sex distribution summaries. These files are generated in the project analysis output directory.")
    doc.add_heading("Acknowledgements", level=1)
    add_paragraph(doc, "To be completed before external submission.")
    doc.add_heading("Authors' contributions", level=1)
    add_paragraph(doc, "To be completed according to author roles before external submission.")
    doc.add_heading("Funding", level=1)
    add_paragraph(doc, "To be completed before external submission.")
    doc.add_heading("Availability of data and materials", level=1)
    add_paragraph(doc, "The Phase 0 audit used CELLxGENE Census release 2025-11-08. Current scripts and summary tables are stored in the local project directory. Before external submission, the analysis code and generated non-sensitive summary tables should be deposited in a public GitHub repository and archived through Zenodo or an equivalent repository to obtain a persistent DOI.")
    doc.add_heading("Ethics approval and consent to participate", level=1)
    add_paragraph(doc, "This study uses public de-identified single-cell metadata and planned public single-cell datasets. No new human participants or animal experiments were included in the current analysis.")
    doc.add_heading("Consent for publication", level=1)
    add_paragraph(doc, "Not applicable.")
    doc.add_heading("Competing interests", level=1)
    add_paragraph(doc, "To be completed before external submission.")

    doc.add_heading("References", level=1)
    refs = [
        "Horvath S. DNA methylation age of human tissues and cell types. Genome Biology. 2013;14:R115. doi:10.1186/gb-2013-14-10-r115.",
        "Hannum G, Guinney J, Zhao L, Zhang L, Hughes G, Sadda S, et al. Genome-wide methylation profiles reveal quantitative views of human aging rates. Molecular Cell. 2013;49:359-367. doi:10.1016/j.molcel.2012.10.016.",
        "Hilton J, et al. CZ CELLxGENE Discover is an online analytical platform and the largest repository of standardized single-cell data. 2024. doi:10.7490/f1000research.1119921.1.",
        "Theodoris CV, Xiao L, Chopra A, Chaffin MD, Al Sayed ZR, Hill MC, et al. Transfer learning enables predictions in network biology. Nature. 2023;618:616-624. doi:10.1038/s41586-023-06139-9.",
        "Cui H, Wang C, Maan H, Pang K, Luo F, Duan N, et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods. 2024. doi:10.1038/s41592-024-02201-0.",
        "Bartz J, Ma X, Zhang L, Dong X. Human Cell Aging Transcriptome Atlas (HCATA): a single-cell atlas of age-associated transcriptomic alterations across human tissues. Communications Biology. 2025. doi:10.1038/s42003-025-08845-8.",
        "Aging Atlas: a multi-omics database for aging biology. Nucleic Acids Research. 2021. doi:10.1093/nar/gkaa894.",
        "The GTEx Consortium atlas of genetic regulatory effects across human tissues. Science. 2020;369:1318-1330. doi:10.1126/science.aaz1776.",
        "Zou Z, Long X, Zhao Q, Zheng Y, Song M, Ma S, et al. A single-cell transcriptomic atlas of human skin aging. Developmental Cell. 2021;56:383-397.e8. doi:10.1016/j.devcel.2020.11.002.",
        "Han X, Zhou Z, Fei L, Sun H, Wang R, Chen Y, et al. Construction of a human cell landscape at single-cell level. Nature. 2020;581:303-309. doi:10.1038/s41586-020-2157-4.",
        "SenNet Portal: Build, Optimization and Usage. 2026. doi:10.64898/2026.02.06.704469.",
        "CELLxGENE Census documentation. CZ CELLxGENE Discover Census and hosted embeddings documentation. Available from: https://chanzuckerberg.github.io/cellxgene-census/.",
    ]
    for ref in refs:
        add_ref(doc, ref)

    for section in doc.sections:
        add_line_numbers(section)

    doc.save(DOCX)
    MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_docx()
